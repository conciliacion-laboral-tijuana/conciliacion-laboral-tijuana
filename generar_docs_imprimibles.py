# -*- coding: utf-8 -*-
"""
Generador de versiones imprimibles de la documentacion del sistema
==================================================================

Convierte los archivos Markdown del proyecto (FEATURES.md, DOCUMENTACION_FUNCIONAL.md,
MANUAL_USUARIO.md, MANUAL_ADMIN.md) en:

  * HTML  (docs/*.html)        - listo para imprimir desde el navegador
  * DOCX  (docs/*.docx)        - editable en Microsoft Word / LibreOffice
  * PDF   (docs/*.pdf)         - via Google Chrome headless (--print-to-pdf)

Requisitos:
  * python-docx (ya incluido en requirements.txt del proyecto)
  * Google Chrome instalado (solo para generar los PDF)

Uso:
    uv run python generar_docs_imprimibles.py
"""

import os
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

PROYECTO = Path(__file__).resolve().parent
SALIDA = PROYECTO / "docs"

DOCUMENTOS = [
    ("FEATURES.md", "System & Features Document (English)"),
    ("DOCUMENTACION_FUNCIONAL.md", "Documentacion Funcional del Sistema"),
    ("MANUAL_USUARIO.md", "Manual de Usuario"),
    ("MANUAL_ADMIN.md", "Manual de Administracion"),
]

CHROME_CANDIDATOS = [
    r"C:\Program Files\Google\Chrome\Application\chrome.exe",
    r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
    os.path.expandvars(r"%LOCALAPPDATA%\Google\Chrome\Application\chrome.exe"),
    os.path.expandvars(r"%PROGRAMFILES%\Google\Chrome\Application\chrome.exe"),
]

VERDE = "1A6B3C"
GRIS_TABLA = "1F2937"


# ═══════════════════════════════════════════════════════════════════════════
#  Parser de Markdown (subconjunto usado en la documentacion del proyecto)
#  Soporta: #/##/###, ---, tablas |...|, listas - y 1., codigo ```, citas >,
#           negrita **, codigo en linea `, enlaces [x](url) y parrafos.
# ═══════════════════════════════════════════════════════════════════════════

_INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def parse_inline(texto):
    """Divide un texto en partes con formato: (texto, tipo[, url])."""
    partes = []
    pos = 0
    for m in _INLINE_RE.finditer(texto):
        if m.start() > pos:
            partes.append((texto[pos:m.start()], "texto"))
        tok = m.group(0)
        if tok.startswith("**"):
            partes.append((tok[2:-2], "negrita"))
        elif tok.startswith("`"):
            partes.append((tok[1:-1], "codigo"))
        else:
            interior = tok[1:-1]
            etiqueta, _, url = interior.partition("](")
            partes.append((etiqueta, "enlace", url))
        pos = m.end()
    if pos < len(texto):
        partes.append((texto[pos:], "texto"))
    return partes


def _es_continuacion_lista(sig):
    """Linea con sangria (2+ espacios) que continua un item de lista."""
    return (
        re.match(r"^\s{2,}\S", sig) is not None
        and re.match(r"^\s*[-*]\s", sig) is None
        and re.match(r"^\s*\d+\.\s", sig) is None
        and not sig.strip().startswith(("|", ">", "```"))
    )


def parse_md(texto):
    """Convierte Markdown en una lista plana de bloques.

    Tipos de bloque:
      ('h1'|'h2'|'h3', texto)
      ('p', texto)
      ('li', texto)              # item de lista con viñeta
      ('ol', (numero, texto))    # item numerado (numero literal)
      ('tabla', (encabezado, filas))
      ('codigo', texto)
      ('cita', texto)
      ('hr', None)
    """
    bloques = []
    lineas = texto.split("\n")
    i, n = 0, len(lineas)

    while i < n:
        linea = lineas[i]

        if not linea.strip():
            i += 1
            continue

        if linea.strip() == "---":
            bloques.append(("hr", None))
            i += 1
            continue

        m = re.match(r"^(#{1,4})\s+(.*)$", linea)
        if m:
            nivel = len(m.group(1))
            bloques.append(("h%d" % min(nivel, 3), m.group(2).strip()))
            i += 1
            continue

        if linea.strip().startswith("```"):
            j = i + 1
            codigo = []
            while j < n and not lineas[j].strip().startswith("```"):
                codigo.append(lineas[j])
                j += 1
            bloques.append(("codigo", "\n".join(codigo)))
            i = j + 1
            continue

        if linea.strip().startswith(">"):
            citas = []
            while i < n and lineas[i].strip().startswith(">"):
                citas.append(lineas[i].strip()[1:].strip())
                i += 1
            bloques.append(("cita", " ".join(citas)))
            continue

        if linea.strip().startswith("|") and i + 1 < n and re.match(r"^\s*\|[\s:\-|]+\|\s*$", lineas[i + 1]):
            encabezado = [c.strip() for c in linea.strip().strip("|").split("|")]
            i += 2
            filas = []
            while i < n and lineas[i].strip().startswith("|"):
                filas.append([c.strip() for c in lineas[i].strip().strip("|").split("|")])
                i += 1
            bloques.append(("tabla", (encabezado, filas)))
            continue

        m = re.match(r"^\s*[-*]\s+(.*)$", linea)
        if m:
            texto_li = m.group(1).strip()
            i += 1
            while i < n and _es_continuacion_lista(lineas[i]):
                texto_li += " " + lineas[i].strip()
                i += 1
            bloques.append(("li", texto_li))
            continue

        m = re.match(r"^\s*(\d+)\.\s+(.*)$", linea)
        if m:
            numero = int(m.group(1))
            texto_ol = m.group(2).strip()
            i += 1
            while i < n and _es_continuacion_lista(lineas[i]):
                texto_ol += " " + lineas[i].strip()
                i += 1
            bloques.append(("ol", (numero, texto_ol)))
            continue

        # Parrafo (une lineas consecutivas no vacias que no inician otro bloque)
        parrafo = [linea.strip()]
        i += 1
        while i < n:
            sig = lineas[i]
            if not sig.strip():
                break
            if (
                re.match(r"^(#{1,4})\s", sig)
                or sig.strip().startswith(("```", ">", "|", "---"))
                or re.match(r"^\s*[-*]\s", sig)
                or re.match(r"^\s*\d+\.\s", sig)
            ):
                break
            parrafo.append(sig.strip())
            i += 1
        bloques.append(("p", " ".join(parrafo)))

    return bloques


# ═══════════════════════════════════════════════════════════════════════════
#  Generacion de HTML (con CSS de impresion A4)
# ═══════════════════════════════════════════════════════════════════════════

_CSS = """
@page { size: A4; margin: 17mm 15mm; }
* { box-sizing: border-box; }
body { font-family: 'Segoe UI', 'Calibri', Arial, sans-serif; font-size: 10.5pt;
       color: #1f2937; line-height: 1.5; margin: 0; }
h1 { font-size: 19pt; color: #111827; border-bottom: 3px solid #1a6b3c;
     padding-bottom: 5px; margin: 4px 0 14px; }
h2 { font-size: 14pt; color: #1a6b3c; margin: 20px 0 8px; page-break-after: avoid; }
h3 { font-size: 11.5pt; color: #374151; margin: 14px 0 6px; page-break-after: avoid; }
p { margin: 6px 0; }
ul, ol { margin: 6px 0 6px 4px; padding-left: 22px; }
li { margin: 2px 0; }
table { border-collapse: collapse; width: 100%; margin: 8px 0; font-size: 9pt;
        page-break-inside: avoid; }
th { background: #1f2937; color: #ffffff; text-align: left; padding: 5px 8px;
     border: 1px solid #374151; }
td { border: 1px solid #d1d5db; padding: 4px 8px; vertical-align: top; }
tr:nth-child(even) td { background: #f9fafb; }
code { font-family: Consolas, 'Courier New', monospace; font-size: 8.8pt;
       background: #eef2ff; padding: 0 3px; border-radius: 3px; }
pre { background: #f3f4f6; border: 1px solid #e5e7eb; padding: 8px 10px;
      border-radius: 4px; white-space: pre-wrap; margin: 8px 0;
      page-break-inside: avoid; }
pre code { background: none; padding: 0; font-size: 8.5pt; }
blockquote { border-left: 4px solid #f59e0b; background: #fffbeb; margin: 8px 0;
             padding: 6px 12px; color: #78350f; }
hr { border: none; border-top: 1px solid #e5e7eb; margin: 16px 0; }
a { color: #2563eb; text-decoration: none; }
strong { color: #111827; }
.cabecera-doc { font-size: 9pt; color: #6b7280; border-bottom: 1px solid #e5e7eb;
                padding-bottom: 6px; margin-bottom: 16px; }
.pie-doc { margin-top: 24px; padding-top: 8px; border-top: 1px solid #e5e7eb;
           font-size: 8.5pt; color: #9ca3af; text-align: center; }
"""


_CSS_COMBINADO = """
section.documento { page-break-before: always; }
.portada { page-break-after: always; height: 250mm; display: flex;
           align-items: center; justify-content: center; }
.portada-inner { text-align: center; width: 100%; }
.portada-logo { font-size: 56pt; line-height: 1.2; }
.portada-titulo { font-size: 34pt; color: #111827; margin: 10px 0 4px; border: none; }
.portada-sub { font-size: 16pt; color: #1a6b3c; font-weight: 600; margin-bottom: 16px; }
.portada-leyenda { font-size: 11pt; color: #4b5563; margin-bottom: 24px; line-height: 1.6; }
.portada-version { display: inline-block; border: 1px solid #d1d5db; border-radius: 999px;
                   padding: 5px 16px; font-size: 9.5pt; color: #374151; margin-bottom: 28px;
                   background: #f9fafb; }
.portada-indice { max-width: 430px; margin: 0 auto; text-align: left;
                  border: 1px solid #e5e7eb; border-radius: 10px; padding: 16px 24px;
                  background: #fafbfc; }
.portada-indice-titulo { font-weight: 700; color: #111827; margin-bottom: 8px; font-size: 11pt; }
.portada-indice ol { margin: 0; padding-left: 22px; font-size: 10.5pt; }
.portada-indice li { margin: 5px 0; }
@page { @bottom-center { content: "Conciliación Laboral Tijuana · Página " counter(page);
         font-size: 8pt; color: #9ca3af; } }
@page :first { @bottom-center { content: none; } }
"""

_MESES_ES = ["enero", "febrero", "marzo", "abril", "mayo", "junio", "julio",
             "agosto", "septiembre", "octubre", "noviembre", "diciembre"]


def portada_html():
    """Portada + índice del documento combinado."""
    from datetime import date
    hoy = date.today()
    fecha = "%d de %s de %d" % (hoy.day, _MESES_ES[hoy.month - 1], hoy.year)
    return (
        '<section class="portada">\n<div class="portada-inner">\n'
        '<div class="portada-logo">⚖️</div>\n'
        '<h1 class="portada-titulo">Despacho Laboral</h1>\n'
        '<div class="portada-sub">Conciliación Laboral Tijuana</div>\n'
        '<div class="portada-leyenda">Documentación completa del sistema:<br>'
        'funcionalidad, manual de usuario, administración y finanzas</div>\n'
        '<div class="portada-version">Versión 1.0 &middot; %s</div>\n'
        '<div class="portada-indice">\n<div class="portada-indice-titulo">Contenido</div>\n<ol>\n'
        '<li>Documentación Funcional del Sistema</li>\n'
        '<li>Manual de Usuario (asesores y administrativos)</li>\n'
        '<li>Manual de Administración (permisos, configuración legal y finanzas)</li>\n'
        '<li>System &amp; Features Document (English)</li>\n'
        '</ol>\n</div>\n</div>\n</section>\n'
    ) % fecha


def esc_html(s):
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def inline_html(partes):
    out = []
    for p in partes:
        if p[1] == "texto":
            out.append(esc_html(p[0]))
        elif p[1] == "negrita":
            out.append("<strong>%s</strong>" % esc_html(p[0]))
        elif p[1] == "codigo":
            out.append("<code>%s</code>" % esc_html(p[0]))
        elif p[1] == "enlace":
            out.append('<a href="%s">%s</a>' % (esc_html(p[2]), esc_html(p[0])))
    return "".join(out)


def bloques_html(bloques):
    partes = []
    i, n = 0, len(bloques)

    def agregar_parrafo(tipo, contenido):
        if tipo == "h1":
            partes.append("<h1>%s</h1>" % esc_html(contenido))
        elif tipo == "h2":
            partes.append("<h2>%s</h2>" % esc_html(contenido))
        elif tipo == "h3":
            partes.append("<h3>%s</h3>" % esc_html(contenido))
        elif tipo == "p":
            partes.append("<p>%s</p>" % inline_html(parse_inline(contenido)))
        elif tipo == "cita":
            partes.append("<blockquote>%s</blockquote>" % inline_html(parse_inline(contenido)))
        elif tipo == "codigo":
            partes.append("<pre><code>%s</code></pre>" % esc_html(contenido))
        elif tipo == "hr":
            partes.append("<hr>")

    while i < n:
        tipo, contenido = bloques[i]

        if tipo == "li":
            items = []
            while i < n and bloques[i][0] == "li":
                items.append(inline_html(parse_inline(bloques[i][1])))
                i += 1
            partes.append("<ul>%s</ul>" % "".join("<li>%s</li>" % it for it in items))
            continue

        if tipo == "ol":
            items = []
            inicio = contenido[0]
            while i < n and bloques[i][0] == "ol":
                items.append(inline_html(parse_inline(bloques[i][1][1])))
                i += 1
            partes.append('<ol start="%d">%s</ol>' % (inicio, "".join("<li>%s</li>" % it for it in items)))
            continue

        if tipo == "tabla":
            encabezado, filas = contenido
            t = ["<table><thead><tr>"]
            t.append("".join("<th>%s</th>" % inline_html(parse_inline(x)) for x in encabezado))
            t.append("</tr></thead><tbody>")
            for fila in filas:
                t.append("<tr>")
                for j, celda in enumerate(fila):
                    t.append("<td>%s</td>" % inline_html(parse_inline(celda)))
                t.append("</tr>")
            t.append("</tbody></table>")
            partes.append("".join(t))
            i += 1
            continue

        agregar_parrafo(tipo, contenido)
        i += 1

    return "\n".join(partes)


def plantilla_html(titulo, cuerpo, css_extra="", con_cabecera=True, con_pie=True):
    cabecera = ('<div class="cabecera-doc">Conciliación Laboral Tijuana &mdash; %s</div>\n'
                % esc_html(titulo)) if con_cabecera else ""
    pie = ('<div class="pie-doc">Generado desde el proyecto Despacho Laboral &middot; '
           "Documento oficial del sistema</div>\n") if con_pie else ""
    return (
        "<!DOCTYPE html>\n<html lang=\"es\">\n<head>\n<meta charset=\"utf-8\">\n"
        "<title>%s</title>\n<style>%s</style>\n</head>\n<body>\n"
        "%s%s%s</body>\n</html>\n"
    ) % (esc_html(titulo), _CSS + css_extra, cabecera, cuerpo, pie)


# ═══════════════════════════════════════════════════════════════════════════
#  Generacion de DOCX con python-docx
# ═══════════════════════════════════════════════════════════════════════════

def _sombreado_celda(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def _sombreado_parrafo(parrafo, color):
    p_pr = parrafo._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    p_pr.append(shd)


def _sombreado_run(run, color):
    r_pr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    r_pr.append(shd)


def _agregar_texto_formateado(parrafo, texto):
    for parte in parse_inline(texto):
        run = parrafo.add_run(parte[0])
        if parte[1] == "negrita":
            run.bold = True
        elif parte[1] == "codigo":
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            _sombreado_run(run, "EEF2FF")
        elif parte[1] == "enlace":
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
            run.underline = True
    return parrafo


def _configurar_pie(doc, titulo):
    sec = doc.sections[0]
    pie = sec.footer.paragraphs[0]
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = pie.add_run("Conciliacion Laboral Tijuana · %s  ·  Pagina " % titulo)
    r1.font.size = Pt(8)
    r1.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    f = pie.add_run()
    f.font.size = Pt(8)
    f.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    f._r.append(fld1)
    f._r.append(instr)
    f._r.append(fld2)


def construir_docx(titulo, bloques):
    doc = Document()

    # Configuracion base: A4 y margenes
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Mm(18)
    sec.bottom_margin = Mm(18)
    sec.left_margin = Mm(20)
    sec.right_margin = Mm(20)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    for tipo, contenido in bloques:
        if tipo == "h1":
            parrafo = doc.add_heading(contenido, level=0)
            for run in parrafo.runs:
                run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
        elif tipo == "h2":
            parrafo = doc.add_heading(contenido, level=1)
            for run in parrafo.runs:
                run.font.color.rgb = RGBColor(0x1A, 0x6B, 0x3C)
        elif tipo == "h3":
            parrafo = doc.add_heading(contenido, level=2)
            for run in parrafo.runs:
                run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
        elif tipo == "p":
            _agregar_texto_formateado(doc.add_paragraph(), contenido)
        elif tipo == "li":
            parrafo = doc.add_paragraph(style="List Bullet")
            _agregar_texto_formateado(parrafo, contenido)
        elif tipo == "ol":
            numero, texto = contenido
            parrafo = doc.add_paragraph()
            parrafo.paragraph_format.left_indent = Mm(6)
            run = parrafo.add_run("%d.  " % numero)
            run.bold = True
            _agregar_texto_formateado(parrafo, texto)
        elif tipo == "cita":
            parrafo = doc.add_paragraph()
            parrafo.paragraph_format.left_indent = Mm(6)
            parrafo.paragraph_format.space_before = Pt(4)
            parrafo.paragraph_format.space_after = Pt(4)
            for run_parte in parse_inline(contenido):
                run = parrafo.add_run(run_parte[0])
                run.italic = True
                if run_parte[1] == "negrita":
                    run.bold = True
        elif tipo == "codigo":
            parrafo = doc.add_paragraph()
            parrafo.paragraph_format.space_before = Pt(4)
            parrafo.paragraph_format.space_after = Pt(4)
            _sombreado_parrafo(parrafo, "F3F4F6")
            lineas = contenido.split("\n")
            for j, linea in enumerate(lineas):
                run = parrafo.add_run(linea if linea else " ")
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)
                if j < len(lineas) - 1:
                    run.add_break()
        elif tipo == "tabla":
            encabezado, filas = contenido
            ncols = len(encabezado)
            tabla = doc.add_table(rows=1 + len(filas), cols=ncols)
            tabla.style = "Table Grid"
            # Encabezado
            for j, celda_txt in enumerate(encabezado):
                cell = tabla.rows[0].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                for run_parte in parse_inline(celda_txt):
                    run = p.add_run(run_parte[0])
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(9)
                _sombreado_celda(cell, GRIS_TABLA)
            # Filas
            for k, fila in enumerate(filas):
                for j in range(ncols):
                    cell = tabla.rows[k + 1].cells[j]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    for run_parte in parse_inline(fila[j] if j < len(fila) else ""):
                        run = p.add_run(run_parte[0])
                        run.font.size = Pt(9)
                        if run_parte[1] == "negrita":
                            run.bold = True
            doc.add_paragraph().paragraph_format.space_after = Pt(0)
        elif tipo == "hr":
            doc.add_paragraph().paragraph_format.space_after = Pt(0)

    _configurar_pie(doc, titulo)
    return doc


# ═══════════════════════════════════════════════════════════════════════════
#  Generacion de PDF con Google Chrome headless
# ═══════════════════════════════════════════════════════════════════════════

def localizar_chrome():
    for ruta in CHROME_CANDIDATOS:
        if Path(ruta).exists():
            return Path(ruta)
    return None


def generar_pdf(chrome, html_ruta, pdf_ruta, user_data):
    url = html_ruta.resolve().as_uri()
    cmd = [
        str(chrome),
        "--headless=new",
        "--disable-gpu",
        "--no-pdf-header-footer",
        "--hide-scrollbars",
        "--run-all-compositor-stages-before-draw",
        "--virtual-time-budget=4000",
        "--user-data-dir=" + str(user_data),
        "--print-to-pdf=" + str(pdf_ruta),
        url,
    ]
    res = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
    if not pdf_ruta.exists() or pdf_ruta.stat().st_size == 0:
        # Reintento con la bandera clasica de Chrome
        cmd2 = [
            str(chrome),
            "--headless",
            "--disable-gpu",
            "--user-data-dir=" + str(user_data),
            "--print-to-pdf=" + str(pdf_ruta),
            url,
        ]
        subprocess.run(cmd2, capture_output=True, text=True, timeout=120)
    return pdf_ruta.exists() and pdf_ruta.stat().st_size > 0


# ═══════════════════════════════════════════════════════════════════════════

def main():
    SALIDA.mkdir(exist_ok=True)
    chrome = localizar_chrome()
    if chrome is None:
        print("AVISO: no se encontro Google Chrome; solo se generaran HTML y DOCX.")

    for nombre, titulo in DOCUMENTOS:
        fuente = PROYECTO / nombre
        if not fuente.exists():
            print("Omitido (no existe):", nombre)
            continue

        md = fuente.read_text(encoding="utf-8")
        bloques = parse_md(md)
        base = nombre.rsplit(".", 1)[0]

        # HTML
        html_ruta = SALIDA / (base + ".html")
        html_ruta.write_text(plantilla_html(titulo, bloques_html(bloques)), encoding="utf-8")
        print("OK HTML:", html_ruta.name)

        # DOCX
        docx_ruta = SALIDA / (base + ".docx")
        construir_docx(titulo, bloques).save(str(docx_ruta))
        print("OK DOCX:", docx_ruta.name, "(%d bytes)" % docx_ruta.stat().st_size)

        # PDF
        if chrome is not None:
            pdf_ruta = SALIDA / (base + ".pdf")
            user_data = SALIDA / (".chrome_tmp_" + base)
            user_data.mkdir(exist_ok=True)
            try:
                if generar_pdf(chrome, html_ruta, pdf_ruta, user_data):
                    print("OK PDF :", pdf_ruta.name, "(%d bytes)" % pdf_ruta.stat().st_size)
                else:
                    print("FALLO PDF:", pdf_ruta.name)
            except Exception as exc:  # noqa: BLE001
                print("ERROR PDF:", pdf_ruta.name, "-", exc)
            finally:
                try:
                    import shutil
                    shutil.rmtree(user_data, ignore_errors=True)
                except Exception:  # noqa: BLE001
                    pass

    # ═══════════════════════════════════════════════════════════════════
    #  Documento combinado: portada + indice + los 4 documentos
    # ═══════════════════════════════════════════════════════════════════
    secciones = [
        ("DOCUMENTACION_FUNCIONAL.md", "Documentación Funcional del Sistema"),
        ("MANUAL_USUARIO.md", "Manual de Usuario"),
        ("MANUAL_ADMIN.md", "Manual de Administración"),
        ("FEATURES.md", "System & Features Document (English)"),
    ]
    cuerpos = []
    for nombre, titulo_seccion in secciones:
        fuente = PROYECTO / nombre
        if not fuente.exists():
            continue
        bloques = parse_md(fuente.read_text(encoding="utf-8"))
        cuerpo = bloques_html(bloques)
        cuerpos.append(
            '<section class="documento">\n'
            '<div class="cabecera-doc">Conciliación Laboral Tijuana &mdash; %s</div>\n'
            "%s\n</section>" % (esc_html(titulo_seccion), cuerpo)
        )

    html_completo = plantilla_html(
        "Documentación Completa del Sistema",
        portada_html() + "\n".join(cuerpos),
        css_extra=_CSS_COMBINADO,
        con_cabecera=False,
        con_pie=False,
    )
    ruta_html = SALIDA / "DOCUMENTACION_COMPLETA.html"
    ruta_html.write_text(html_completo, encoding="utf-8")
    print("OK HTML:", ruta_html.name)

    if chrome is not None:
        pdf_ruta = SALIDA / "DOCUMENTACION_COMPLETA.pdf"
        user_data = SALIDA / ".chrome_tmp_completa"
        user_data.mkdir(exist_ok=True)
        try:
            if generar_pdf(chrome, ruta_html, pdf_ruta, user_data):
                print("OK PDF :", pdf_ruta.name, "(%d bytes)" % pdf_ruta.stat().st_size)
            else:
                print("FALLO PDF:", pdf_ruta.name)
        except Exception as exc:  # noqa: BLE001
            print("ERROR PDF:", pdf_ruta.name, "-", exc)
        finally:
            import shutil
            shutil.rmtree(user_data, ignore_errors=True)

    print("\nListo. Archivos generados en:", SALIDA)


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:  # noqa: BLE001
        print("ERROR GENERAL:", exc, file=sys.stderr)
        sys.exit(1)
