"""
Tests automatizados: verificación de cálculos laborales en demandas.
=====================================================================

Ejecuta el comando seed_clientes_prueba (crea casos con datos completos
por tipo de despido) y después la verificación completa de cálculos
(verify_demandas_calculos.verificar_calculos). Si el cálculo o la
generación de demandas se rompe, el test falla.

También valida el flujo completo vía la app: cliente con datos reales →
cálculo laboral (GET/POST) → descarga de la demanda Word.

Uso:
    uv run python manage.py test expedientes.tests
"""
import json
import sys
import unicodedata
from datetime import date, timedelta
from decimal import Decimal
from django.utils import timezone
from io import BytesIO, StringIO


def _normalizar(texto):
    """Normaliza texto: mayúsculas y sin acentos (para comparar contenido de docx)."""
    if not texto:
        return ''
    return unicodedata.normalize('NFD', texto.upper()).encode('ascii', 'ignore').decode('ascii')

from django.contrib.auth.models import User
from django.core.management import call_command
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import Client, TestCase
from django.urls import reverse

from expedientes.models import Cliente, Expediente, CalculoLaboral, TareaConciliacion, Aviso


def _silenciar_call_command():
    """Redirige stdout para silenciar la salida del comando."""
    old_stdout = sys.stdout
    sys.stdout = StringIO()
    return old_stdout


def _restaurar_stdout(old_stdout):
    sys.stdout = old_stdout


class CalculosDemandaTests(TestCase):
    """Verifica que los cálculos de las demandas sean completos y consistentes."""

    @classmethod
    def setUpTestData(cls):
        # Asegurar que exista al menos un asesor (requisito del comando)
        if not User.objects.filter(profile__rol='asesor').exists():
            asesor = User.objects.create_user(
                username='asesor_test_ci', password='x',
                first_name='Asesor', last_name='CI',
            )
            asesor.profile.rol = 'asesor'
            asesor.profile.save()

        # Sembrar los casos de prueba con datos completos
        old_stdout = _silenciar_call_command()
        try:
            call_command('seed_clientes_prueba')
        finally:
            _restaurar_stdout(old_stdout)

    def test_casos_de_prueba_creados(self):
        """El comando seed_clientes_prueba debe crear los 9 casos marcados."""
        from expedientes.management.commands.seed_clientes_prueba import NOMBRES_PRUEBA
        self.assertEqual(
            Cliente.objects.filter(nombre__in=NOMBRES_PRUEBA).count(),
            len(NOMBRES_PRUEBA),
            'Deben existir los 9 clientes de prueba con datos completos',
        )
        self.assertGreaterEqual(Expediente.objects.count(), 5)
        self.assertGreaterEqual(CalculoLaboral.objects.count(), 5)

    def test_verificacion_calculos_sin_problemas(self):
        """La verificación completa de cálculos no debe detectar problemas."""
        from verify_demandas_calculos import verificar_calculos

        resultado = verificar_calculos()

        # Al menos los casos de prueba deben tener datos completos
        self.assertGreaterEqual(
            resultado['con_datos'],
            len(Expediente.objects.filter(estado='demanda')),
            'Todos los casos de prueba en demanda deben tener datos para calcular',
        )
        # Sin problemas en ninguna verificación (conceptos, consistencia, totales)
        self.assertEqual(
            resultado['problemas'], [],
            f'Se detectaron problemas en los cálculos: {resultado["problemas"]}',
        )

    def test_renuncia_voluntaria_no_reclama_indemnizacion(self):
        """En renuncia voluntaria la demanda NO debe reclamar indemnización 90d ni prima antigüedad."""
        from verify_demandas_calculos import verificar_calculos

        resultado = verificar_calculos()

        voluntarios = [r for r in resultado['resumen'] if r['tipo'] == 'voluntario']
        self.assertGreaterEqual(len(voluntarios), 2, 'Deben existir casos voluntarios de prueba')

        # Ningún problema de tipo 'RENUNCIA VOLUNTARIA pero la demanda reclama...'
        problemas_voluntario = [
            p for p in resultado['problemas']
            if 'RENUNCIA VOLUNTARIA' in p
        ]
        self.assertEqual(
            problemas_voluntario, [],
            f'Las renuncias voluntarias reclaman conceptos que no proceden: {problemas_voluntario}',
        )

    def test_totales_coinciden_con_calculo_laboral_guardado(self):
        """El total de la demanda debe coincidir con el CalculoLaboral guardado."""
        from verify_demandas_calculos import verificar_calculos

        resultado = verificar_calculos()

        # Para cada caso en demanda, el total de la tabla (o del cálculo) debe
        # coincidir con el CalculoLaboral guardado
        discrepancias = [
            p for p in resultado['problemas']
            if 'CalculoLaboral.total' in p
        ]
        self.assertEqual(
            discrepancias, [],
            f'Discrepancias entre CalculoLaboral y demanda: {discrepancias}',
        )


class FlujoCompletoDemandaTests(TestCase):
    """
    Valida el flujo completo vía la app (con cliente autenticado):

        cliente con datos reales → cálculo laboral (GET crea el cálculo,
        POST lo recalcula) → descarga de la demanda Word.

    Simula exactamente lo que hace una abogada en el sistema:
    1. Abre la página de cálculo laboral del expediente
    2. Recalcula con el formulario
    3. Descarga la demanda Word (vista 'generar_demanda')
    """

    @classmethod
    def setUpTestData(cls):
        # Usuarios: asesor asignado al caso + abogada que genera documentos
        cls.asesor = User.objects.create_user(
            username='asesor_flow_test', password='x',
            first_name='Asesor', last_name='Flow',
        )
        cls.asesor.profile.rol = 'asesor'
        cls.asesor.profile.save()

        cls.abogada = User.objects.create_user(
            username='abogada_flow_test', password='x',
            first_name='Abogada', last_name='Flow',
        )
        cls.abogada.profile.rol = 'abogada'
        cls.abogada.profile.save()

        # Cliente con DATOS REALES completos (fechas + salario) — requisito
        # para que el cálculo no quede vacío
        hoy = date.today()
        cls.cliente = Cliente.objects.create(
            nombre='Cliente Flujo Prueba',
            curp='CURP12345678901234',
            telefono='+526641234567',
            whatsapp='+526641234567',
            direccion_calle='Av. Revolución',
            direccion_numero='1234',
            direccion_cp='22000',
            direccion_colonia='Zona Centro',
            empresa='Empresa Flujo SA de CV',
            empresa_telefono='+526649990000',
            empresa_calle='Blvd. Industrial',
            empresa_numero='500',
            empresa_colonia='Otay',
            empresa_cp='22400',
            puesto='Operador de producción',
            salario=Decimal('15000.00'),
            periodo_pago='mensual',
            horas_semanales=48,
            jornada='diurna',
            fecha_ingreso=hoy - timedelta(days=365 * 3 + 30),
            fecha_salida=hoy - timedelta(days=15),
            genero='masculino',
            oficina='plaza_patria',
            como_supo='google',
        )

        cls.expediente = Expediente.objects.create(
            cliente=cls.cliente,
            asesor=cls.asesor,
            estado='demanda',
            tipo_despido='injustificado',
            prioridad='alta',
            notas='Caso creado para test de flujo completo.',
        )

        # Segundo caso: RENUNCIA VOLUNTARIA (mismo cliente, datos completos)
        cls.cliente_voluntario = Cliente.objects.create(
            nombre='Cliente Renuncia Voluntaria',
            telefono='+526642345678',
            whatsapp='+526642345678',
            direccion_calle='Calle Renuncia',
            direccion_numero='77',
            direccion_cp='22000',
            direccion_colonia='Centro',
            empresa='Empresa Renuncia SA',
            empresa_telefono='+526649991111',
            puesto='Auxiliar administrativo',
            salario=Decimal('12000.00'),
            periodo_pago='mensual',
            horas_semanales=40,
            jornada='diurna',
            fecha_ingreso=hoy - timedelta(days=365 * 2 + 10),
            fecha_salida=hoy - timedelta(days=7),
            genero='femenino',
            oficina='otay',
            como_supo='recomendacion',
        )
        cls.expediente_voluntario = Expediente.objects.create(
            cliente=cls.cliente_voluntario,
            asesor=cls.asesor,
            estado='demanda',
            tipo_despido='voluntario',
            prioridad='media',
            notas='Renuncia voluntaria: solo prestaciones proporcionales.',
        )

    def setUp(self):
        self.client = Client()
        self.assertTrue(
            self.client.login(username='abogada_flow_test', password='x'),
            'La abogada de prueba debe poder iniciar sesión',
        )

    def _urls(self):
        return {
            'calculo': reverse('calculo_laboral', args=[self.expediente.pk]),
            'demanda': reverse('generar_demanda', args=[self.expediente.pk]),
        }

    def test_get_calculo_laboral_crea_calculo_con_total(self):
        """GET a cálculo laboral crea el CalculoLaboral y calcula el total."""
        urls = self._urls()

        resp = self.client.get(urls['calculo'])
        self.assertEqual(resp.status_code, 200, 'La página de cálculo debe cargar')

        calculo = CalculoLaboral.objects.get(expediente=self.expediente)
        self.assertGreater(
            calculo.total, Decimal('0'),
            'El cálculo con datos reales debe producir un total > 0',
        )
        self.assertGreater(calculo.dias_trabajados, 0)
        # El total aparece en la página
        self.assertContains(resp, f'${calculo.total:,.2f}')

    def test_post_calculo_laboral_recalcula(self):
        """POST del formulario recalcula y guarda el total (flujo real de la vista)."""
        urls = self._urls()

        datos_post = {
            'periodo_pago': 'mensual',
            'notas': 'Recálculo desde el test de flujo',
            # Checkboxes de conceptos (todos los base activos en injustificado)
            'incluir_aguinaldo': 'on',
            'incluir_vacaciones': 'on',
            'incluir_prima_vacacional': 'on',
            'incluir_prima_antiguedad': 'on',
            'incluir_indemnizacion': 'on',
            'incluir_indemnizacion_20dias': '',
            'incluir_vacaciones_vencidas': '',
            'incluir_horas_extras': '',
            'incluir_salarios_devengados': '',
            'incluir_dias_festivos': '',
            # Campos de entrada
            'dias_vacaciones_vencidos': '0',
            'horas_extra_cantidad': '0',
            'salarios_devengados': '0',
            'dias_festivos_cantidad': '0',
            'dias_vacaciones_override': '',
        }

        resp = self.client.post(urls['calculo'], datos_post)
        self.assertEqual(resp.status_code, 302, 'Tras guardar debe redirigir al cálculo')
        self.assertEqual(resp.url, urls['calculo'])

        calculo = CalculoLaboral.objects.get(expediente=self.expediente)
        self.assertGreater(
            calculo.total, Decimal('0'),
            'El recálculo con datos reales debe producir un total > 0',
        )
        self.assertTrue(calculo.incluir_indemnizacion)
        self.assertTrue(calculo.incluir_prima_antiguedad)

    def test_descarga_demanda_word_contiene_datos_y_total(self):
        """Descarga de la demanda (vista generar_demanda) devuelve un .docx válido."""
        urls = self._urls()

        # El cálculo debe existir antes de generar la demanda
        resp = self.client.get(urls['calculo'])
        self.assertEqual(resp.status_code, 200)
        calculo = CalculoLaboral.objects.get(expediente=self.expediente)
        self.assertGreater(calculo.total, 0)

        # Descarga directa de la demanda
        resp = self.client.get(urls['demanda'])
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(
            resp['Content-Type'],
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        self.assertIn('attachment; filename=', resp['Content-Disposition'])
        self.assertIn('Demanda_', resp['Content-Disposition'])
        self.assertGreater(len(resp.content), 5000, 'El .docx no puede estar vacío')

        # Leer el .docx generado y verificar contenido
        from docx import Document
        doc = Document(BytesIO(resp.content))
        texto = '\n'.join(p.text for p in doc.paragraphs)
        tabla = '\n'.join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )
        contenido = (texto + '\n' + tabla).upper()

        self.assertIn('CLIENTE FLUJO PRUEBA', contenido, 'La demanda debe contener al cliente')
        self.assertIn('PRESTACIONES', contenido, 'La demanda debe tener la sección de prestaciones')
        self.assertIn('TOTAL', contenido, 'La demanda debe tener la fila TOTAL')

        # El total de la tabla de prestaciones coincide numéricamente con el
        # CalculoLaboral guardado (el .docx lo muestra como "$106,143.31" en
        # la fila ['', 'TOTAL:', '$106,143.31'])
        fila_total = None
        for table in doc.tables:
            for row in table.rows:
                if any('TOTAL' in c.text.upper() for c in row.cells[:2]):
                    fila_total = row
                    break
            if fila_total:
                break
        self.assertIsNotNone(
            fila_total,
            'La tabla de prestaciones del .docx debe contener la fila TOTAL',
        )
        total_tabla = fila_total.cells[-1].text.strip()
        monto = Decimal(total_tabla.replace('$', '').replace(',', '').strip())
        self.assertEqual(
            monto, calculo.total,
            f'El total de la tabla ({total_tabla}) debe coincidir con el CalculoLaboral ({calculo.total})',
        )

        # Movimiento registrado al descargar
        self.assertTrue(
            self.expediente.movimientos.filter(
                detalle__icontains='Documento de demanda laboral generado'
            ).exists(),
            'La descarga debe registrar un movimiento',
        )

    def test_guard_descarga_bloquea_sin_datos_criticos(self):
        """El guard de datos críticos bloquea la descarga si falta CURP/salario/fechas."""
        # Cliente sin salario ni fechas (datos incompletos)
        cliente_incompleto = Cliente.objects.create(
            nombre='Cliente Incompleto Guard',
            oficina='plaza_patria',
        )
        exp_incompleto = Expediente.objects.create(
            cliente=cliente_incompleto,
            asesor=self.asesor,
            estado='demanda',
            tipo_despido='injustificado',
        )

        resp = self.client.post(
            reverse('demanda_descargar', args=[exp_incompleto.pk]),
            {'contenido': '<p>Demanda</p>'},
        )
        self.assertEqual(resp.status_code, 302)
        self.assertIn('asistente', resp.url, 'Debe redirigir al asistente para completar datos')

    def test_wizard_completo_genera_demanda_y_guarda_machote(self):
        """El asistente paso a paso avanza, valida datos críticos y guarda machotes con marcadores."""
        url = reverse('demanda_asistente', args=[self.expediente.pk])

        # El expediente base tiene datos completos, pero probamos que el flujo avance
        # paso 1 → 2 → 3 → 4 y termine en el editor
        resp = self.client.post(url, {
            'paso_actual': '1', 'accion': 'siguiente',
            'nombre': self.cliente.nombre,
            'curp': self.cliente.curp,
            'oficina': self.cliente.oficina,
            'genero': self.cliente.genero,
        })
        self.assertEqual(resp.status_code, 302)
        self.assertIn('paso=2', resp.url, 'El asistente debe avanzar al paso 2')

        resp = self.client.post(url, {
            'paso_actual': '2', 'accion': 'siguiente',
            'salario': str(self.cliente.salario),
            'fecha_ingreso': self.cliente.fecha_ingreso.isoformat(),
            'fecha_salida': self.cliente.fecha_salida.isoformat(),
        })
        self.assertIn('paso=3', resp.url, 'El asistente debe avanzar al paso 3')

        resp = self.client.post(url, {
            'paso_actual': '3', 'accion': 'siguiente',
            'empresa': self.cliente.empresa,
        })
        self.assertIn('paso=4', resp.url, 'El asistente debe avanzar al paso 4 (revisión)')

        # Paso 4: revisión con cálculo y firma
        resp = self.client.get(f'{url}?paso=4')
        self.assertEqual(resp.status_code, 200)
        body = resp.content.decode('utf-8', errors='replace')
        self.assertIn('tipo_despido', body, 'El paso 4 debe ofrecer tipo de despido')
        self.assertIn('TOTAL', body, 'El paso 4 debe mostrar el cálculo')

        # Finalizar → editor de demanda
        resp = self.client.post(url, {
            'paso_actual': '4', 'accion': 'finalizar', 'tipo_despido': 'injustificado',
        })
        self.assertEqual(resp.status_code, 302)
        self.assertIn('demanda', resp.url)

        # Guardar como machote (con nombre único del test)
        resp = self.client.post(reverse('demanda_guardar_machote', args=[self.expediente.pk]), {
            'nombre': 'Machote Test Wizard',
            'contenido': f'<p>{self.cliente.nombre}</p><p>Salario: $15,000.00 mensuales</p>'
                         f'<p>{self.cliente.empresa}</p>',
        })
        self.assertEqual(resp.status_code, 302)
        from expedientes.models import Machote
        machote = Machote.objects.filter(nombre='Machote Test Wizard').first()
        self.assertIsNotNone(machote, 'El machote debe crearse')
        self.assertIn('{{ nombre_cliente }}', machote.contenido_html, 'El nombre debe convertirse en marcador')
        self.assertIn('{{ nombre_empresa }}', machote.contenido_html, 'La empresa debe convertirse en marcador')

    def test_editar_y_eliminar_machote_desde_catalogo(self):
        """El catálogo permite editar y eliminar machotes."""
        from expedientes.models import Machote

        machote = Machote.objects.create(
            nombre='Machote Test Edicion',
            descripcion='Plantilla de prueba',
            categoria='demanda',
            contenido_html='<p>{{ nombre_cliente }} original</p>',
            icono='📄',
            activo=True,
        )

        # Editar: GET muestra el contenido y POST guarda los cambios
        resp = self.client.get(reverse('machote_editar', args=[machote.pk]))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(resp, 'original')

        resp = self.client.post(reverse('machote_editar', args=[machote.pk]), {
            'nombre': 'Machote Test Editado',
            'descripcion': 'Descripcion nueva',
            'categoria': 'demanda',
            'tipo_despido': 'injustificado',
            'jurisdiccion': 'federal',
            'icono': '⚡',
            'contenido': '<p>{{ nombre_cliente }} actualizado</p>',
        })
        self.assertEqual(resp.status_code, 302)
        machote.refresh_from_db()
        self.assertEqual(machote.nombre, 'Machote Test Editado')
        self.assertEqual(machote.tipo_despido, 'injustificado')
        self.assertIn('actualizado', machote.contenido_html)

        # Catálogo: muestra botones Editar y Eliminar
        resp = self.client.get(reverse('machotes_catalogo'))
        body = resp.content.decode('utf-8', errors='replace')
        self.assertIn('Editar', body)
        self.assertIn('Eliminar', body)

        # Eliminar: POST lo borra y el detalle devuelve 404
        resp = self.client.post(reverse('machote_eliminar', args=[machote.pk]))
        self.assertEqual(resp.status_code, 302)
        self.assertFalse(Machote.objects.filter(pk=machote.pk).exists())
        resp = self.client.get(reverse('machote_editar', args=[machote.pk]))
        self.assertEqual(resp.status_code, 404)

    def test_renombrar_machote_desde_catalogo(self):
        """La abogada puede renombrar un machote desde el catálogo sin tocar el contenido."""
        from expedientes.models import Machote

        machote = Machote.objects.create(
            nombre='Machote Nombre Viejo',
            descripcion='Plantilla de prueba',
            categoria='demanda',
            contenido_html='<p>{{ nombre_cliente }} del caso</p>',
            icono='📄',
            activo=True,
        )

        # El catálogo muestra la acción Renombrar (botón con data-url) para la abogada
        resp = self.client.get(reverse('machotes_catalogo'))
        self.assertEqual(resp.status_code, 200)
        self.assertContains(
            resp,
            reverse('machote_renombrar', args=[machote.pk]),
            msg_prefix='El catálogo debe renderizar el botón Renombrar con su URL',
        )

        # POST renombra: solo cambia el nombre, el contenido HTML no se altera
        resp = self.client.post(reverse('machote_renombrar', args=[machote.pk]), {
            'nombre': 'Machote Nombre Nuevo',
        })
        self.assertEqual(resp.status_code, 302)
        self.assertEqual(resp.url, reverse('machotes_catalogo'))

        machote.refresh_from_db()
        self.assertEqual(machote.nombre, 'Machote Nombre Nuevo')
        self.assertIn(
            'nombre_cliente', machote.contenido_html,
            'Al renombrar no debe alterarse el contenido de la plantilla',
        )

        # Nombre vacío o solo espacios se rechaza sin cambiar nada
        resp = self.client.post(reverse('machote_renombrar', args=[machote.pk]), {'nombre': '   '})
        self.assertEqual(resp.status_code, 302)
        machote.refresh_from_db()
        self.assertEqual(machote.nombre, 'Machote Nombre Nuevo')

    def test_modal_renombrar_compartido_en_catalogo_y_dashboard(self):
        """El modal compartido (partial) aparece tanto en el catálogo como en el dashboard de la abogada."""
        from expedientes.models import Machote

        machote = Machote.objects.create(
            nombre='Machote Modal Compartido',
            descripcion='Plantilla de prueba',
            categoria='demanda',
            contenido_html='<p>{{ nombre_cliente }} del caso</p>',
            icono='📄',
            activo=True,
        )

        # ── Catálogo de machotes ──
        resp = self.client.get(reverse('machotes_catalogo'))
        self.assertEqual(resp.status_code, 200)
        body = resp.content.decode('utf-8', errors='replace')
        self.assertIn('modal-renombrar', body, 'El catálogo debe incluir el modal compartido')
        self.assertIn('abrirRenombrar', body, 'El catálogo debe incluir el JS del modal')
        self.assertIn(
            f'name="next" value="{reverse("machotes_catalogo")}"',
            body,
            'El modal del catálogo debe regresar al propio catálogo (request.path)',
        )

        # ── Dashboard de la abogada ──
        resp = self.client.get(reverse('dashboard_abogada'))
        self.assertEqual(resp.status_code, 200)
        body = resp.content.decode('utf-8', errors='replace')
        self.assertIn('modal-renombrar', body, 'El dashboard debe incluir el modal compartido')
        self.assertIn('abrirRenombrar', body, 'El dashboard debe incluir el JS del modal')
        self.assertContains(
            resp,
            reverse('machote_renombrar', args=[machote.pk]),
            msg_prefix='El dashboard debe renderizar el botón Renombrar en la fila del machote',
        )
        self.assertIn(
            f'name="next" value="{reverse("dashboard_abogada")}"',
            body,
            'El modal del dashboard debe regresar al propio dashboard (request.path)',
        )

    def test_renombrar_machote_regresa_a_pagina_origen(self):
        """Renombrar desde el catálogo vuelve al catálogo y desde el dashboard vuelve al dashboard."""
        from expedientes.models import Machote

        machote = Machote.objects.create(
            nombre='Machote Origen',
            descripcion='Plantilla de prueba',
            categoria='demanda',
            contenido_html='<p>Contenido sin marcadores</p>',
            icono='📄',
            activo=True,
        )

        # ── Desde el catálogo: next = ruta del catálogo ──
        resp = self.client.post(reverse('machote_renombrar', args=[machote.pk]), {
            'nombre': 'Machote Origen Renombrado',
            'next': reverse('machotes_catalogo'),
        })
        self.assertEqual(resp.status_code, 302)
        self.assertEqual(
            resp.url, reverse('machotes_catalogo'),
            'Al renombrar desde el catálogo debe volver al catálogo',
        )
        machote.refresh_from_db()
        self.assertEqual(machote.nombre, 'Machote Origen Renombrado')

        # ── Desde el dashboard de la abogada: next = ruta del dashboard ──
        resp = self.client.post(reverse('machote_renombrar', args=[machote.pk]), {
            'nombre': 'Machote Origen Final',
            'next': reverse('dashboard_abogada'),
        })
        self.assertEqual(resp.status_code, 302)
        self.assertEqual(
            resp.url, reverse('dashboard_abogada'),
            'Al renombrar desde el dashboard debe volver al dashboard de la abogada',
        )
        machote.refresh_from_db()
        self.assertEqual(machote.nombre, 'Machote Origen Final')

        # ── next inseguro (URL absoluta externa): cae al catálogo por defecto ──
        # (se envía un nombre distinto: el renombrado sí ocurre, pero el redirect
        #  debe ser al catálogo y NO a la URL externa — sin open redirect)
        resp = self.client.post(reverse('machote_renombrar', args=[machote.pk]), {
            'nombre': 'Machote Hackeado',
            'next': 'https://malicious.example/phishing',
        })
        self.assertEqual(resp.status_code, 302)
        self.assertEqual(
            resp.url, reverse('machotes_catalogo'),
            'Un next inseguro debe caer al catálogo por defecto (sin open redirect)',
        )
        machote.refresh_from_db()
        self.assertEqual(
            machote.nombre, 'Machote Hackeado',
            'El renombrado debe haberse aplicado aunque el next sea inseguro',
        )

    def test_listado_expedientes_muestra_y_filtra_por_oficina(self):
        """El listado de expedientes muestra la oficina del caso y permite filtrar por ella."""
        # Las únicas oficinas válidas son Plaza Patria, Otay y CLT
        self.assertEqual(
            [k for k, _ in Cliente.OFICINA_CHOICES],
            ['plaza_patria', 'otay', 'clt'],
            'Las únicas oficinas deben ser Plaza Patria, Otay y CLT',
        )

        # Crear un caso en oficina CLT (los existentes son plaza_patria y otay)
        cliente_clt = Cliente.objects.create(
            nombre='Cliente CLT Filtro', oficina='clt', empresa='Empresa CLT',
        )
        Expediente.objects.create(cliente=cliente_clt, asesor=self.asesor, estado='nuevo')

        # El listado muestra la columna Oficina y el badge del caso CLT
        resp = self.client.get(reverse('expediente_list'))
        self.assertEqual(resp.status_code, 200)
        body = resp.content.decode('utf-8', errors='replace')
        self.assertIn('>Oficina<', body, 'El listado debe tener la columna Oficina')
        self.assertIn('>CLT<', body, 'El listado debe mostrar el badge de la oficina CLT')

        # Filtrar por oficina=clt devuelve solo los casos de esa oficina
        resp = self.client.get(reverse('expediente_list'), {'oficina': 'clt'})
        body = resp.content.decode('utf-8', errors='replace')
        self.assertIn('Cliente CLT Filtro', body, 'El filtro debe incluir el caso CLT')

        # El detalle del expediente muestra la oficina
        resp = self.client.get(reverse('expediente_detail', args=[
            Expediente.objects.get(cliente=cliente_clt).pk,
        ]))
        body = resp.content.decode('utf-8', errors='replace')
        self.assertIn('Oficina', body)
        self.assertIn('CLT', body)

    def test_reemplazar_marcadores_no_crashea_con_datos_faltantes(self):
        """El reemplazo de marcadores no crashea cuando un campo del cliente es None."""
        from expedientes.marcadores import reemplazar_marcadores

        # Cliente incompleto: sin CURP, sin teléfono, sin salario ni fechas
        cliente_incompleto = Cliente.objects.create(
            nombre='Cliente Incompleto Marcadores',
            oficina='plaza_patria',
        )
        exp = Expediente.objects.create(
            cliente=cliente_incompleto,
            asesor=self.asesor,
            estado='demanda',
            tipo_despido='injustificado',
        )

        html = (
            '<p>{{ nombre_cliente }} | {{ curp }} | {{ telefono }} | {{ rfc }} | '
            '{{ numero_expediente }} | {{ tipo_despido }}</p>'
        )
        # Antes del fix esto lanzaba TypeError: replace() argument 2 must be str, not None
        resultado = reemplazar_marcadores(html, exp)

        self.assertIn(exp.numero, resultado, 'El número de expediente debe reemplazarse')
        self.assertIn(
            exp.get_tipo_despido_display(), resultado,
            'El tipo de despido debe reemplazarse',
        )
        # Los campos faltantes muestran placeholders (no espacios en blanco)
        self.assertIn('[CURP]', resultado, 'El CURP faltante debe mostrar el placeholder [CURP]')
        self.assertIn('[TELÉFONO]', resultado, 'El teléfono faltante debe mostrar el placeholder [TELÉFONO]')
        self.assertIn(
            '[RFC]', resultado,
            'Los marcadores con fallback de placeholder deben conservar su comportamiento',
        )
        self.assertNotIn('None', resultado, 'Ningún marcador debe quedar como el texto "None"')
        self.assertNotIn('{{', resultado, 'No deben quedar marcadores sin reemplazar')

    def test_renombrar_machote_requiere_permiso(self):
        """Un asesor sin permiso de documentos no puede renombrar machotes."""
        from expedientes.models import Machote

        machote = Machote.objects.create(
            nombre='Machote Protegido',
            categoria='demanda',
            contenido_html='<p>Contenido</p>',
            icono='📄',
            activo=True,
        )

        # El asesor (sin permiso de documentos) es rechazado y no cambia nada
        self.assertTrue(self.client.login(username='asesor_flow_test', password='x'))
        resp = self.client.post(reverse('machote_renombrar', args=[machote.pk]), {
            'nombre': 'Renombrado por asesor',
        })
        self.assertEqual(resp.status_code, 302)
        self.assertEqual(resp.url, reverse('machotes_catalogo'))

        machote.refresh_from_db()
        self.assertEqual(
            machote.nombre, 'Machote Protegido',
            'El asesor sin permiso no debe poder renombrar machotes',
        )

    def test_catalogo_busqueda_y_filtro_por_categoria(self):
        """El catálogo permite buscar por texto y filtrar por categoría, combinados."""
        from expedientes.models import Machote

        Machote.objects.create(
            nombre='Demanda por Despido Injustificado',
            descripcion='Plantilla para casos de despido',
            categoria='demanda',
            contenido_html='<p>Demanda laboral</p>',
            icono='⚡', activo=True,
        )
        Machote.objects.create(
            nombre='Carta Finiquito Básica',
            descripcion='Finiquito de renuncia',
            categoria='carta_finiquito',
            contenido_html='<p>Carta finiquito</p>',
            icono='📄', activo=True,
        )
        Machote.objects.create(
            nombre='Convenio de Pagos',
            descripcion='Convenio extrajudicial',
            categoria='convenio',
            contenido_html='<p>Convenio</p>',
            icono='🤝', activo=True,
        )

        url_catalogo = reverse('machotes_catalogo')

        # 1) Búsqueda por texto: solo debe aparecer el que coincide
        resp = self.client.get(url_catalogo, {'q': 'finiquito'})
        body = resp.content.decode('utf-8', errors='replace')
        self.assertIn('Carta Finiquito Básica', body)
        self.assertNotIn('Demanda por Despido Injustificado', body)
        self.assertNotIn('Convenio de Pagos', body)
        # Indica cuántos resultados hay para la búsqueda
        self.assertIn('1 resultado', body)

        # 2) Filtro por categoría
        resp = self.client.get(url_catalogo, {'categoria': 'demanda'})
        body = resp.content.decode('utf-8', errors='replace')
        self.assertIn('Demanda por Despido Injustificado', body)
        self.assertNotIn('Carta Finiquito Básica', body)
        self.assertNotIn('Convenio de Pagos', body)

        # 3) Combinación: búsqueda + categoría que no coincide → vacío
        resp = self.client.get(url_catalogo, {'q': 'despido', 'categoria': 'convenio'})
        body = resp.content.decode('utf-8', errors='replace')
        self.assertNotIn('Demanda por Despido Injustificado', body)
        self.assertNotIn('Convenio de Pagos', body)
        self.assertIn('No hay machotes que coincidan', body)

        # 4) Sin filtros: se muestran todos
        resp = self.client.get(url_catalogo)
        body = resp.content.decode('utf-8', errors='replace')
        self.assertIn('Demanda por Despido Injustificado', body)
        self.assertIn('Carta Finiquito Básica', body)
        self.assertIn('Convenio de Pagos', body)
        self.assertIn('3 plantillas', body)

    def test_renuncia_voluntaria_no_reclama_indemnizacion_ni_prima_antiguedad(self):
        """La demanda de renuncia voluntaria NO debe incluir indemnización (Art. 50) ni prima de antigüedad (Art. 162)."""
        # 1) El cálculo laboral se crea y los conceptos quedan desactivados
        resp = self.client.get(reverse('calculo_laboral', args=[self.expediente_voluntario.pk]))
        self.assertEqual(resp.status_code, 200)

        calculo = CalculoLaboral.objects.get(expediente=self.expediente_voluntario)
        self.assertGreater(calculo.total, 0, 'El cálculo de renuncia debe tener total > 0')
        self.assertFalse(
            calculo.incluir_indemnizacion,
            'Renuncia voluntaria: la indemnización constitucional debe estar desactivada',
        )
        self.assertFalse(
            calculo.incluir_prima_antiguedad,
            'Renuncia voluntaria: la prima de antigüedad debe estar desactivada',
        )
        self.assertEqual(
            calculo.indemnizacion, Decimal('0'),
            'El monto de indemnización debe ser 0 en renuncia voluntaria',
        )
        self.assertEqual(
            calculo.prima_antiguedad, Decimal('0'),
            'El monto de prima de antigüedad debe ser 0 en renuncia voluntaria',
        )

        # 2) La demanda descargada no debe incluir esos conceptos
        resp = self.client.get(reverse('generar_demanda', args=[self.expediente_voluntario.pk]))
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(
            resp['Content-Type'],
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )

        from docx import Document
        doc = Document(BytesIO(resp.content))

        # Encontrar la tabla de prestaciones (la que tiene encabezado PRESTACIÓN)
        # Nota: se normaliza sin acentos porque el .docx los guarda como caracteres
        # compuestos y el texto normalizado evita falsos negativos de comparación.
        tabla_prestaciones = None
        for table in doc.tables:
            if any('PRESTACION' in _normalizar(cell.text) for row in table.rows for cell in row.cells):
                tabla_prestaciones = table
                break
        self.assertIsNotNone(
            tabla_prestaciones,
            'La demanda debe contener la tabla de prestaciones',
        )
        # Solo se comparan las columnas de concepto/fundamento (nunca los montos,
        # para que '50'/'162' de una cantidad no generen falsos positivos)
        etiquetas_filas = [_normalizar(row.cells[0].text) for row in tabla_prestaciones.rows]
        fundamentos = [_normalizar(row.cells[1].text) for row in tabla_prestaciones.rows]

        # Las prestaciones proporcionales SÍ se reclaman
        self.assertIn('AGUINALDO PROPORCIONAL', etiquetas_filas)
        self.assertIn('VACACIONES PROPORCIONALES', etiquetas_filas)
        self.assertIn('PRIMA VACACIONAL (25%)', etiquetas_filas)
        self.assertIn('TOTAL', ' '.join(fundamentos), 'La tabla debe cerrar con TOTAL')

        # NO se reclama prima de antigüedad ni indemnización (por concepto ni artículo)
        self.assertNotIn(
            'PRIMA DE ANTIGUEDAD', etiquetas_filas,
            'Renuncia voluntaria NO debe reclamar prima de antigüedad',
        )
        self.assertNotIn(
            'INDEMNIZACION CONSTITUCIONAL', etiquetas_filas,
            'Renuncia voluntaria NO debe reclamar indemnización',
        )
        self.assertNotIn(
            'ART. 162', ' '.join(fundamentos),
            'No debe citarse el Art. 162 LFT (prima de antigüedad)',
        )
        self.assertNotIn(
            'ART. 50', ' '.join(fundamentos),
            'No debe citarse el Art. 50 LFT (indemnización 3 meses)',
        )

        # El total de la tabla sigue siendo el del cálculo
        fila_total = next(
            row for row in tabla_prestaciones.rows
            if any('TOTAL' in _normalizar(c.text) for c in row.cells[:2])
        )
        monto_total = Decimal(fila_total.cells[-1].text.strip().replace('$', '').replace(',', ''))
        self.assertEqual(monto_total, calculo.total)


class ImportarEmpresasTests(TestCase):
    """Prueba el parser de domicilios y la importación de Empresas y Domicilios.xlsx."""

    def test_desglosar_domicilio_centro_tijuana(self):
        from expedientes.management.commands.importar_empresas import _desglosar_domicilio
        calle, numero, colonia, cp = _desglosar_domicilio(
            'AVENIDA MELCHOR OCAMPO NO. 729, CENTRO TIJUANA, BAJA CALIFORNIA'
        )
        self.assertEqual(calle, 'MELCHOR OCAMPO')
        self.assertEqual(numero, '729')
        self.assertEqual(colonia, 'CENTRO')
        self.assertEqual(cp, '')

    def test_desglosar_domicilio_no_confunde_numero_con_cp(self):
        from expedientes.management.commands.importar_empresas import _desglosar_domicilio
        _, numero, colonia, cp = _desglosar_domicilio(
            'BLVD GUADALAJARA NO. 15510, LA MESA, TIJUANA BC.'
        )
        self.assertEqual(numero, '15510')
        self.assertEqual(colonia, 'LA MESA')
        self.assertEqual(cp, '', 'El número de la calle no debe detectarse como CP')

    def test_desglosar_domicilio_numero_compuesto(self):
        from expedientes.management.commands.importar_empresas import _desglosar_domicilio
        _, numero, _, _ = _desglosar_domicilio(
            'CALLE SANTA ROSALIA NO. 20386-20302, BUENOS AIRES SUR, TIJUANA, BAJA CALIFORNIA'
        )
        self.assertEqual(numero, '20386-20302')

    def test_desglosar_domicilio_cp_explicito(self):
        from expedientes.management.commands.importar_empresas import _desglosar_domicilio
        calle, numero, colonia, cp = _desglosar_domicilio(
            'AVENIDA DE LA AMISTAD,NO.S/N, EMPLEADOS FEDERALES,TIJUANA, BAJA CALIFORNIA, CP. 22010.'
        )
        self.assertEqual(calle, 'DE LA AMISTAD')
        self.assertEqual(numero, '')
        self.assertEqual(colonia, 'EMPLEADOS FEDERALES')
        self.assertEqual(cp, '22010')

    def test_desglosar_domicilio_tijuana_baja_california_separados(self):
        from expedientes.management.commands.importar_empresas import _desglosar_domicilio
        _, _, colonia, _ = _desglosar_domicilio(
            'CALLA REAL DE MAESTRANZA NO. 7361, REAL DE SAN FRANCISCO, TIJUANA, BAJA CALIFORNIA.'
        )
        self.assertEqual(colonia, 'REAL DE SAN FRANCISCO')

    def test_desglosar_domicilio_vacio(self):
        from expedientes.management.commands.importar_empresas import _desglosar_domicilio
        self.assertEqual(_desglosar_domicilio(''), ('', '', '', ''))
        self.assertEqual(_desglosar_domicilio(None), ('', '', '', ''))

    def test_detectar_tipo_persona(self):
        from expedientes.management.commands.importar_empresas import _detectar_tipo_persona
        self.assertEqual(_detectar_tipo_persona('MISSION FOODS MEXICO S DE RL DE CV'), 'moral')
        self.assertEqual(_detectar_tipo_persona('COMERCIALIZADORA VEBRA, S.A. DE C.V.'), 'moral')
        self.assertEqual(_detectar_tipo_persona('SECRETARIA DE DEFENSA NACIONAL'), 'moral')
        self.assertEqual(_detectar_tipo_persona('INSTITUTO FRONTERA A.C.'), 'moral')
        self.assertEqual(_detectar_tipo_persona('MARIA CRISTINA LOZA DE LA TORRE'), 'fisica')
        self.assertEqual(_detectar_tipo_persona('LEONARDO SANDOVAL ESTRADA'), 'fisica')

    def test_importacion_idempotente(self):
        import os
        import tempfile

        from django.core.management import call_command
        from openpyxl import Workbook

        from expedientes.models import Empresa

        with tempfile.TemporaryDirectory() as tmp:
            ruta = os.path.join(tmp, 'empresas.xlsx')
            wb = Workbook()
            ws = wb.active
            ws.append(['EMPRESA', 'DOMICILIO', 'ABOGADO', 'TELEFONO'])
            ws.append(['PRUEBA SA DE CV', 'CALLE EJEMPLO NO. 10, COLONIA CENTRO, TIJUANA, BAJA CALIFORNIA', 'LIC JUAN PEREZ', '664-123-4567'])
            ws.append(['MARIA GOMEZ LOPEZ', 'AV PRINCIPAL NO. 20, ZONA RIO, TIJUANA, BAJA CALIFORNIA C.P. 22000', '', ''])
            wb.save(ruta)

            call_command('importar_empresas', archivo=ruta, verbosity=0)
            self.assertEqual(Empresa.objects.count(), 2)

            # Re-ejecutar no debe duplicar
            call_command('importar_empresas', archivo=ruta, verbosity=0)
            self.assertEqual(Empresa.objects.count(), 2)

            emp = Empresa.objects.get(nombre='PRUEBA SA DE CV')
            self.assertEqual(emp.tipo_persona, 'moral')
            self.assertEqual(emp.domicilio_calle, 'EJEMPLO')
            self.assertEqual(emp.domicilio_numero, '10')
            self.assertEqual(emp.abogado, 'LIC JUAN PEREZ')

            fisica = Empresa.objects.get(nombre='MARIA GOMEZ LOPEZ')
            self.assertEqual(fisica.tipo_persona, 'fisica')
            self.assertEqual(fisica.domicilio_cp, '22000')


class AcuseParserTests(TestCase):
    """Prueba el parser del acuse de solicitud de conciliación."""

    @classmethod
    def setUpTestData(cls):
        cls.asesor = User.objects.create_user(
            username='asesor_acuse_parser', password='x',
            first_name='Asesor', last_name='Parser',
        )
        cls.asesor.profile.rol = 'asesor'
        cls.asesor.profile.save()

    TEXTO_ACUSE = """
ACUSE DE SOLICITUD DE CONCILIACIÓN
FECHA DE SOLICITUD: 07 de Agosto de 2026
SOLICITANTE(S): JOSE LIMON DIAZ
CITADO(S): TAQUERIA LOS ALBAÑILES
FECHA DE CONFLICTO: 30 de Julio de 2026
OBJETO DE LA CONCILIACIÓN: Despido
UNIDAD DE CONCILIACIÓN TIJUANA
Usted ha guardado exitosamente la solicitud de conciliación con folio TIJ/26427/2026.
"""

    def test_parsear_acuse_completo(self):
        from expedientes.acuse_parser import parsear_acuse

        datos = parsear_acuse(self.TEXTO_ACUSE)
        self.assertEqual(datos['folio'], 'TIJ/26427/2026')
        self.assertEqual(datos['solicitante'], 'JOSE LIMON DIAZ')
        self.assertEqual(datos['citado'], 'TAQUERIA LOS ALBAÑILES')
        self.assertEqual(datos['fecha_solicitud'], date(2026, 8, 7))
        self.assertEqual(datos['fecha_conflicto'], date(2026, 7, 30))
        self.assertEqual(datos['objeto'], 'Despido')
        self.assertEqual(datos['tipo_despido'], 'injustificado')
        self.assertEqual(datos['unidad'], 'TIJUANA')

    def test_objeto_terminacion_voluntaria(self):
        from expedientes.acuse_parser import parsear_acuse

        texto = self.TEXTO_ACUSE.replace('Despido', 'Terminación voluntaria')
        datos = parsear_acuse(texto)
        self.assertEqual(datos['tipo_despido'], 'voluntario')

    def test_objeto_rescision(self):
        from expedientes.acuse_parser import parsear_acuse

        texto = self.TEXTO_ACUSE.replace('Despido', 'Rescisión')
        datos = parsear_acuse(texto)
        self.assertEqual(datos['tipo_despido'], 'rescision')

    def test_parsear_fecha_es(self):
        from expedientes.acuse_parser import _parsear_fecha_es

        self.assertEqual(_parsear_fecha_es('07 de Agosto de 2026'), date(2026, 8, 7))
        self.assertEqual(_parsear_fecha_es('1 de Enero de 2025'), date(2025, 1, 1))
        self.assertIsNone(_parsear_fecha_es('no es una fecha'))

    def test_parsear_pdf_real(self):
        """El parser debe leer el acuse real (si está presente en el repo)."""
        from pathlib import Path
        from expedientes.acuse_parser import parsear_acuse_pdf

        ruta = Path('acuse-jose-diaz.pdf')
        if not ruta.exists():
            self.skipTest('acuse-jose-diaz.pdf no está presente')

        datos = parsear_acuse_pdf(ruta.read_bytes())
        self.assertEqual(datos.get('folio'), 'TIJ/26427/2026')
        self.assertEqual(datos.get('solicitante'), 'JOSE LIMON DIAZ')
        self.assertEqual(datos.get('citado'), 'TAQUERIA LOS ALBAÑILES')
        self.assertEqual(datos.get('tipo_despido'), 'injustificado')

    def test_mapear_campos_modelo(self):
        from expedientes.acuse_parser import mapear_campos_modelo, parsear_acuse

        cliente = Cliente.objects.create(nombre='Cliente Viejo', oficina='plaza_patria')
        expediente = Expediente.objects.create(
            cliente=cliente,
            asesor=self.asesor,
            estado='nuevo',
            tipo_despido='injustificado',
        )
        datos = parsear_acuse(self.TEXTO_ACUSE)
        campos = mapear_campos_modelo(datos, expediente)

        por_key = {c['key']: c for c in campos}
        self.assertIn('folio', por_key)
        self.assertIn('nombre', por_key)
        self.assertIn('citado', por_key)
        self.assertEqual(por_key['nombre']['valor_detectado'], 'JOSE LIMON DIAZ')
        # El cliente tiene nombre ('Cliente Viejo') → el campo NO es nuevo pero difiere
        self.assertFalse(por_key['nombre']['nuevo'])
        self.assertTrue(por_key['nombre']['difiere'], 'El nombre distinto debe marcarse como difiere')
        # El tipo de despido ya está puesto → no es nuevo ni difiere
        self.assertFalse(por_key['tipo_despido']['nuevo'])
        self.assertFalse(por_key['tipo_despido']['difiere'])
        # La unidad se incluye con el valor detectado (se crea la solicitud al confirmar)
        self.assertIn('unidad', por_key)
        self.assertEqual(por_key['unidad']['valor_detectado'], 'TIJUANA')
        self.assertTrue(por_key['unidad']['nuevo'])


class FlujoAcuseTests(TestCase):
    """Prueba el flujo: subir acuse → vista previa → confirmar datos."""

    def setUp(self):
        self.asesor = User.objects.create_user(
            username='asesor_acuse', password='x', first_name='Asesor', last_name='Acuse',
        )
        self.asesor.profile.rol = 'asesor'
        self.asesor.profile.save()

        self.cliente = Cliente.objects.create(
            nombre='JOSE LIMON DIAZ',
            curp='CURP12345678901234',
            telefono='+526641234567',
            direccion_calle='Av. Revolución',
            direccion_numero='1234',
            direccion_cp='22000',
            direccion_colonia='Zona Centro',
            empresa='TAQUERIA LOS ALBAÑILES',
            oficina='plaza_patria',
        )
        self.expediente = Expediente.objects.create(
            cliente=self.cliente,
            asesor=self.asesor,
            estado='nuevo',
            tipo_despido='injustificado',
        )
        self.client.login(username='asesor_acuse', password='x')

    def _pdf_bytes(self):
        """Genera un PDF mínimo válido con el texto del acuse dentro."""
        import fitz
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), 'ACUSE DE SOLICITUD DE CONCILIACIÓN')
        page.insert_text((72, 90), 'FECHA DE SOLICITUD: 07 de Agosto de 2026')
        page.insert_text((72, 108), 'SOLICITANTE(S): JOSE LIMON DIAZ')
        page.insert_text((72, 126), 'CITADO(S): TAQUERIA LOS ALBAÑILES')
        page.insert_text((72, 144), 'FECHA DE CONFLICTO: 30 de Julio de 2026')
        page.insert_text((72, 162), 'OBJETO DE LA CONCILIACIÓN: Despido')
        page.insert_text((72, 180), 'UNIDAD DE CONCILIACIÓN TIJUANA')
        page.insert_text((72, 198), 'con folio TIJ/26427/2026')
        return doc.tobytes()

    def test_subir_pdf_redirige_a_vista_previa(self):
        url = reverse('subir_conciliacion_pdf', args=[self.expediente.pk])
        resp = self.client.post(url, {'pdf': SimpleUploadedFile('acuse.pdf', self._pdf_bytes(), 'application/pdf')})
        self.assertEqual(resp.status_code, 302)
        self.assertIn('acuse-vista-previa', resp.url)

    def test_vista_previa_muestra_campos_detectados(self):
        url = reverse('subir_conciliacion_pdf', args=[self.expediente.pk])
        self.client.post(url, {'pdf': SimpleUploadedFile('acuse.pdf', self._pdf_bytes(), 'application/pdf')})

        resp = self.client.get(reverse('acuse_vista_previa', args=[self.expediente.pk]))
        self.assertEqual(resp.status_code, 200)
        body = resp.content.decode('utf-8', errors='replace')
        self.assertIn('TIJ/26427/2026', body, 'El folio detectado debe mostrarse')
        self.assertIn('JOSE LIMON DIAZ', body, 'El solicitante debe mostrarse')
        self.assertIn('TAQUERIA LOS ALBAÑILES', body, 'El citado debe mostrarse')
        self.assertIn('name="campos"', body, 'Debe haber checkboxes de campos')

    def test_confirmar_aplica_folio_y_fecha_tramite(self):
        url = reverse('subir_conciliacion_pdf', args=[self.expediente.pk])
        self.client.post(url, {'pdf': SimpleUploadedFile('acuse.pdf', self._pdf_bytes(), 'application/pdf')})

        resp = self.client.post(reverse('confirmar_acuse_datos', args=[self.expediente.pk]), {
            'campos': ['folio', 'fecha_solicitud', 'nombre', 'citado', 'fecha_conflicto', 'tipo_despido'],
        })
        self.assertEqual(resp.status_code, 302)

        self.expediente.refresh_from_db()
        self.cliente.refresh_from_db()
        self.assertEqual(self.expediente.folio, 'TIJ/26427/2026')
        self.assertEqual(self.expediente.fecha_tramite, date(2026, 8, 7))
        self.assertEqual(self.cliente.nombre, 'JOSE LIMON DIAZ')
        self.assertEqual(self.cliente.empresa, 'TAQUERIA LOS ALBAÑILES')
        self.assertEqual(self.cliente.fecha_salida, date(2026, 7, 30))
        self.assertEqual(self.expediente.tipo_despido, 'injustificado')
        # El estado avanza de nuevo → solicitud
        self.assertEqual(self.expediente.estado, 'solicitud')
        # Se guardó el documento
        self.assertTrue(self.expediente.documentos.filter(descripcion__icontains='Acuse de Conciliación').exists())
        # Tarea de conciliación completada
        self.assertTrue(self.expediente.tareas_conciliacion.filter(estado='completado').exists())

    def test_confirmar_sin_seleccion_no_aplica(self):
        url = reverse('subir_conciliacion_pdf', args=[self.expediente.pk])
        self.client.post(url, {'pdf': SimpleUploadedFile('acuse.pdf', self._pdf_bytes(), 'application/pdf')})

        resp = self.client.post(reverse('confirmar_acuse_datos', args=[self.expediente.pk]), {'campos': []})
        self.assertEqual(resp.status_code, 302)
        self.expediente.refresh_from_db()
        self.assertEqual(self.expediente.folio, '', 'Sin selección no debe aplicarse el folio')

    def test_guardar_pdf_sin_datos_detectados(self):
        """PDF sin datos parseables aún puede guardarse como documento (modo solo_guardar)."""
        url = reverse('subir_conciliacion_pdf', args=[self.expediente.pk])
        # PDF mínimo sin texto reconocible del acuse
        resp = self.client.post(url, {'pdf': SimpleUploadedFile('otro.pdf', b'%PDF-1.4\n%\xe2\xe3\xcf\xd3', 'application/pdf')})
        self.assertEqual(resp.status_code, 302)

        resp = self.client.post(reverse('confirmar_acuse_datos', args=[self.expediente.pk]), {
            'accion': 'solo_guardar',
        })
        self.assertEqual(resp.status_code, 302)
        self.assertTrue(
            self.expediente.documentos.filter(descripcion__icontains='Acuse de Conciliación').exists(),
            'El PDF debe guardarse como documento aunque no tenga datos',
        )

    def test_confirmar_requiere_post(self):
        """GET a confirmar_acuse_datos debe redirigir sin aplicar nada."""
        url = reverse('subir_conciliacion_pdf', args=[self.expediente.pk])
        self.client.post(url, {'pdf': SimpleUploadedFile('acuse.pdf', self._pdf_bytes(), 'application/pdf')})

        resp = self.client.get(reverse('confirmar_acuse_datos', args=[self.expediente.pk]))
        self.assertEqual(resp.status_code, 405, 'GET debe rechazarse (requiere POST)')

    def test_sin_pdf_pendiente_redirige_al_detalle(self):
        resp = self.client.get(reverse('acuse_vista_previa', args=[self.expediente.pk]))
        self.assertEqual(resp.status_code, 302)
        self.assertIn('expedientes', resp.url)

    def test_rechaza_no_pdf(self):
        url = reverse('subir_conciliacion_pdf', args=[self.expediente.pk])
        resp = self.client.post(url, {'pdf': SimpleUploadedFile('acuse.txt', b'no es pdf', 'text/plain')})
        self.assertEqual(resp.status_code, 302)
        self.assertIn('expedientes', resp.url)


class EspejoEnVivoTests(TestCase):
    """Prueba el espejo en vivo: screenshots servibles y endpoint de estado."""

    def setUp(self):
        self.asesor = User.objects.create_user(
            username='asesor_espejo', password='x', first_name='Asesor', last_name='Espejo',
        )
        self.asesor.profile.rol = 'asesor'
        self.asesor.profile.save()

        self.cliente = Cliente.objects.create(
            nombre='Cliente Espejo', curp='CURP12345678901234',
            telefono='+526641234567', oficina='plaza_patria',
        )
        self.expediente = Expediente.objects.create(
            cliente=self.cliente, asesor=self.asesor, estado='nuevo',
        )
        self.client.login(username='asesor_espejo', password='x')

    def _crear_tarea_con_screenshots(self, estado='ejecutando', urls=None):
        from expedientes.models import TareaConciliacion
        tarea = TareaConciliacion.objects.create(
            expediente=self.expediente, usuario=self.asesor, estado=estado,
        )
        if urls:
            tarea.screenshots_json = json.dumps(urls)
            tarea.save(update_fields=['screenshots_json'])
        return tarea

    def test_screenshots_a_urls_convierte_rutas_media(self):
        from expedientes.conciliacion_automation import screenshots_a_urls
        from django.conf import settings
        from pathlib import Path

        media = Path(settings.MEDIA_ROOT)
        ruta = str(media / 'conciliacion' / 'tarea_1' / '04_solicitante.png')
        urls = screenshots_a_urls([ruta])
        self.assertEqual(urls, ['/media/conciliacion/tarea_1/04_solicitante.png'])

    def test_screenshots_a_urls_ignora_rutas_fuera_de_media(self):
        from expedientes.conciliacion_automation import screenshots_a_urls
        urls = screenshots_a_urls(['C:\\temp\\fuera.png', ''])
        self.assertEqual(urls, ['', ''])

    def test_estado_devuelve_screenshots_guardados(self):
        """Las URLs guardadas se convierten a la vista autenticada (no /media/ público)."""
        tarea = self._crear_tarea_con_screenshots(
            estado='completado',
            urls=['/media/conciliacion/tarea_1/07_enviado.png'],
        )
        resp = self.client.get(reverse('conciliacion_estado', args=[tarea.pk]))
        self.assertEqual(resp.status_code, 200)
        data = resp.json()
        self.assertEqual(data['estado'], 'completado')
        # La URL debe apuntar a la vista autenticada con el nombre del archivo
        self.assertEqual(len(data['screenshots']), 1)
        url = data['screenshots'][0]
        self.assertIn(f'/conciliacion/{tarea.pk}/screenshot/', url)
        self.assertTrue(url.endswith('07_enviado.png'))

    def test_estado_escanea_directorio_durante_ejecucion(self):
        """Durante la ejecución, el endpoint descubre capturas parciales del directorio."""
        import shutil
        from expedientes.views import _directorio_screenshots_tarea

        tarea = self._crear_tarea_con_screenshots(estado='ejecutando', urls=None)

        # Limpiar directorio compartido de MEDIA (los PKs se reusan entre clases
        # de test, pero el MEDIA_ROOT es global y puede tener archivos residuales)
        directorio = _directorio_screenshots_tarea(tarea.pk)
        if directorio.exists():
            shutil.rmtree(directorio)

        # Crear capturas parciales en el directorio de la tarea
        directorio.mkdir(parents=True, exist_ok=True)
        (directorio / '00_inicio.png').write_bytes(b'x')
        (directorio / '04_solicitante.png').write_bytes(b'x')

        resp = self.client.get(reverse('conciliacion_estado', args=[tarea.pk]))
        data = resp.json()
        self.assertEqual(data['estado'], 'ejecutando')
        self.assertEqual(len(data['screenshots']), 2)
        self.assertTrue(data['screenshots'][0].endswith('00_inicio.png'))
        self.assertIn(f'/conciliacion/{tarea.pk}/screenshot/', data['screenshots'][0])

    def test_estado_sin_screenshots_devuelve_vacio(self):
        import shutil
        from expedientes.views import _directorio_screenshots_tarea

        tarea = self._crear_tarea_con_screenshots(estado='ejecutando', urls=None)
        # Asegurar que el directorio de esta tarea no exista (sin capturas parciales)
        directorio = _directorio_screenshots_tarea(tarea.pk)
        if directorio.exists():
            shutil.rmtree(directorio)
        resp = self.client.get(reverse('conciliacion_estado', args=[tarea.pk]))
        data = resp.json()
        self.assertEqual(data['screenshots'], [])

    def test_screenshot_vista_requiere_login(self):
        """La vista de screenshot exige autenticación (protege PII de clientes)."""
        from django.contrib.auth.models import User as DjangoUser
        from django.contrib.auth import logout
        from expedientes.views import _directorio_screenshots_tarea

        tarea = self._crear_tarea_con_screenshots(estado='ejecutando')
        directorio = _directorio_screenshots_tarea(tarea.pk)
        directorio.mkdir(parents=True, exist_ok=True)
        (directorio / '00_inicio.png').write_bytes(b'PNGDATA')

        # Cerrar sesión (setUp hizo login)
        self.client.logout()
        resp = self.client.get(reverse('conciliacion_screenshot', args=[tarea.pk, '00_inicio.png']))
        # Sin login → redirect al login
        self.assertEqual(resp.status_code, 302)

    def test_screenshot_vista_sirve_archivo_con_acceso(self):
        tarea = self._crear_tarea_con_screenshots(estado='ejecutando')
        from expedientes.views import _directorio_screenshots_tarea

        directorio = _directorio_screenshots_tarea(tarea.pk)
        directorio.mkdir(parents=True, exist_ok=True)
        (directorio / '00_inicio.png').write_bytes(b'PNGDATA')

        # setUp ya hizo login con asesor_espejo (acceso al expediente)
        resp = self.client.get(reverse('conciliacion_screenshot', args=[tarea.pk, '00_inicio.png']))
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(b''.join(resp.streaming_content), b'PNGDATA')

    def test_screenshot_vista_bloquea_usuario_sin_acceso(self):
        """Un usuario sin acceso al expediente no puede ver los screenshots."""
        from django.contrib.auth.models import User as DjangoUser
        from expedientes.views import _directorio_screenshots_tarea

        tarea = self._crear_tarea_con_screenshots(estado='ejecutando')
        directorio = _directorio_screenshots_tarea(tarea.pk)
        directorio.mkdir(parents=True, exist_ok=True)
        (directorio / '00_inicio.png').write_bytes(b'PNGDATA')

        otro = DjangoUser.objects.create_user(username='otrouser', password='clave123')
        self.client.logout()
        self.client.login(username='otrouser', password='clave123')
        resp = self.client.get(reverse('conciliacion_screenshot', args=[tarea.pk, '00_inicio.png']))
        self.assertEqual(resp.status_code, 404)

    def test_screenshot_vista_rechaza_path_traversal(self):
        """Un nombre con traversal no puede leer archivos fuera del directorio de la tarea."""
        from expedientes.views import _directorio_screenshots_tarea

        tarea = self._crear_tarea_con_screenshots(estado='ejecutando')
        directorio = _directorio_screenshots_tarea(tarea.pk)
        directorio.mkdir(parents=True, exist_ok=True)
        (directorio / '00_inicio.png').write_bytes(b'PNGDATA')

        # La vista sanitiza con Path(...).name: '..%2F..%2Fsecret.png' → nombre basename
        # que no existe en el directorio de la tarea → 404 (no se filtra nada)
        resp = self.client.get(f'/conciliacion/{tarea.pk}/screenshot/..%2F..%2Fsecret.png')
        self.assertEqual(resp.status_code, 404)

    def test_procesando_contiene_espejo(self):
        tarea = self._crear_tarea_con_screenshots(estado='ejecutando')
        resp = self.client.get(reverse('conciliacion_procesando', args=[tarea.pk]))
        self.assertEqual(resp.status_code, 200)
        body = resp.content.decode('utf-8', errors='replace')
        self.assertIn('mirror-panel', body, 'La página debe incluir el panel de espejo en vivo')
        self.assertIn('mirror-img', body, 'La página debe incluir la imagen del espejo')
        self.assertIn('espejo', body.lower())


class BuscarEmpresasCatalogoTests(TestCase):
    """Prueba el endpoint de autocompletado del catálogo de empresas."""

    def setUp(self):
        from expedientes.models import Empresa
        self.user = User.objects.create_user(username='usuarioprueba', password='clave123')
        Empresa.objects.create(
            nombre='MISION FOODS MEXICO S DE RL DE CV',
            domicilio='AVENIDA MELCHOR OCAMPO NO. 729, CENTRO TIJUANA, BAJA CALIFORNIA',
            domicilio_calle='MELCHOR OCAMPO',
            domicilio_numero='729',
            domicilio_colonia='CENTRO',
            tipo_persona='moral',
        )

    def test_busqueda_devuelve_empresa_con_datos(self):
        self.client.login(username='usuarioprueba', password='clave123')
        r = self.client.get(reverse('buscar_empresas_catalogo'), {'q': 'MISION'})
        self.assertEqual(r.status_code, 200)
        data = r.json()
        self.assertEqual(len(data['resultados']), 1)
        e = data['resultados'][0]
        self.assertEqual(e['nombre'], 'MISION FOODS MEXICO S DE RL DE CV')
        self.assertEqual(e['calle'], 'MELCHOR OCAMPO')
        self.assertEqual(e['numero'], '729')
        self.assertEqual(e['tipo_persona'], 'moral')

    def test_busqueda_requiere_login(self):
        r = self.client.get(reverse('buscar_empresas_catalogo'), {'q': 'MISION'})
        self.assertEqual(r.status_code, 302)

    def test_busqueda_exige_minimo_2_caracteres(self):
        self.client.login(username='usuarioprueba', password='clave123')
        r = self.client.get(reverse('buscar_empresas_catalogo'), {'q': 'M'})
        self.assertEqual(r.status_code, 200)
        self.assertEqual(r.json()['resultados'], [])


class ExtensionChromeApiTests(TestCase):
    """
    Tests de la API REST que usa la Extensión de Chrome:

    - GET  /api/extension/tareas/           → tareas pendientes con datos para llenar
    - POST /api/extension/tareas/<id>/reportar/ → folio + acuse + screenshots
    - Autenticación por token (Authorization: Token <token>)
    - Página de configuración de la extensión (token + instrucciones)
    """

    @classmethod
    def setUpTestData(cls):
        from expedientes.models import Empresa

        # Asesor con token
        cls.asesor = User.objects.create_user(
            username='ext_asesor', password='x',
            first_name='Ext', last_name='Asesor',
        )
        cls.asesor.profile.rol = 'asesor'
        cls.asesor.profile.save()
        cls.token = cls.asesor.profile.api_token

        # Otro asesor (sin acceso al expediente)
        cls.otro = User.objects.create_user(
            username='ext_otro', password='x', first_name='Otro', last_name='Asesor',
        )
        cls.otro.profile.rol = 'asesor'
        cls.otro.profile.save()

        cls.cliente = Cliente.objects.create(
            nombre='Juan Carlos López Moreno',
            curp='LOOM800101HTCPBN07',
            rfc='LOMA800101ABC',
            telefono='6641234567',
            email='juan.lopez@email.com',
            genero='masculino',
            fecha_nacimiento=date(1980, 1, 1),
            direccion_calle='Calle Uno',
            direccion_numero='123',
            direccion_cp='22000',
            empresa='Mi Empresa SA de CV',
            empresa_razon_social='Mi Empresa SA de CV',
            tipo_persona_citado='moral',
            puesto='Operador',
            salario=Decimal('12000.00'),
            periodo_pago='mensual',
            jornada='diurna',
            fecha_ingreso=date(2022, 3, 1),
            fecha_salida=date(2025, 6, 30),
            oficina='plaza_patria',
        )
        cls.expediente = Expediente.objects.create(
            cliente=cls.cliente,
            asesor=cls.asesor,
            estado='nuevo',
        )

        cls.tarea = TareaConciliacion.objects.create(
            expediente=cls.expediente,
            usuario=cls.asesor,
            estado='pendiente',
            modo='extension',
        )

    def _auth(self, token=None):
        return {'HTTP_AUTHORIZATION': f'Token {token or self.token}'}

    # ─── GET /api/extension/tareas/ ───────────────────────────────────

    def test_tareas_sin_token_es_401(self):
        r = self.client.get(reverse('extension_api_tareas'))
        self.assertEqual(r.status_code, 401)

    def test_tareas_con_token_invalido_es_401(self):
        r = self.client.get(reverse('extension_api_tareas'), **self._auth('token-malo'))
        self.assertEqual(r.status_code, 401)

    def test_tareas_devuelve_datos_para_llenar(self):
        r = self.client.get(reverse('extension_api_tareas'), **self._auth())
        self.assertEqual(r.status_code, 200)
        data = r.json()
        self.assertEqual(len(data['tareas']), 1)
        t = data['tareas'][0]
        self.assertEqual(t['id'], self.tarea.pk)
        self.assertEqual(t['expediente']['numero'], self.expediente.numero)
        self.assertEqual(t['cliente']['nombre'], 'Juan Carlos López Moreno')
        self.assertEqual(t['cliente']['curp'], 'LOOM800101HTCPBN07')
        self.assertEqual(t['cliente']['tipo_persona'], '2')          # moral
        self.assertEqual(t['cliente']['genero'], '1')                # masculino
        self.assertEqual(t['cliente']['periodicidad'], '2')          # mensual
        self.assertEqual(t['cliente']['fecha_ingreso'], '01/03/2022')
        self.assertEqual(t['cliente']['fecha_salida'], '30/06/2025')
        # Persona Moral: razón social y datos del citado
        self.assertEqual(t['cliente']['empresa_nombre'], 'Mi Empresa SA de CV')
        self.assertEqual(t['cliente']['empresa_rfc'], 'LOMA800101ABC')
        self.assertEqual(t['cliente']['empresa_email'], 'juan.lopez@email.com')
        self.assertTrue(t['portal']['url_solicitud'].startswith('https://app.conciliacionbc'))
        self.assertIn('fui despedido', t['hechos'])

    def test_tareas_persona_fisica_campos_corresponden(self):
        """Un cliente persona física debe enviar tipo_persona='1' y los campos del citado correctos."""
        cliente_fisica = Cliente.objects.create(
            nombre='Maria Elena Garcia Ruiz',
            curp='GARR850310MTCRRL09',
            rfc='GARR850310A99',
            telefono='6649876543',
            email='maria.garcia@email.com',
            genero='femenino',
            fecha_nacimiento=date(1985, 3, 10),
            direccion_calle='Av. Constitucion',
            direccion_numero='456',
            direccion_cp='22000',
            empresa='Juan Perez Lopez',
            empresa_razon_social='Juan Perez Lopez',
            tipo_persona_citado='fisica',
            empresa_telefono='6645551234',
            empresa_calle='Calle Falsa',
            empresa_numero='789',
            empresa_cp='22300',
            puesto='Cajera',
            salario=Decimal('8000.00'),
            periodo_pago='semanal',
            jornada='diurna',
            fecha_ingreso=date(2023, 6, 1),
            fecha_salida=date(2025, 7, 15),
            oficina='otay',
        )
        exp_fisica = Expediente.objects.create(
            cliente=cliente_fisica, asesor=self.asesor, estado='nuevo',
        )
        tarea_fisica = TareaConciliacion.objects.create(
            expediente=exp_fisica, usuario=self.asesor,
            estado='pendiente', modo='extension',
        )

        r = self.client.get(reverse('extension_api_tareas'), **self._auth())
        self.assertEqual(r.status_code, 200)
        tareas = r.json()['tareas']
        t_fisica = next(t for t in tareas if t['id'] == tarea_fisica.pk)

        # Persona Física: tipo_persona='1'
        self.assertEqual(t_fisica['cliente']['tipo_persona'], '1')
        self.assertEqual(t_fisica['cliente']['empresa_nombre'], 'Juan Perez Lopez')
        self.assertEqual(t_fisica['cliente']['empresa_rfc'], 'GARR850310A99')
        self.assertEqual(t_fisica['cliente']['empresa_email'], 'maria.garcia@email.com')
        self.assertEqual(t_fisica['cliente']['empresa_telefono'], '6645551234')

        # Limpiar
        tarea_fisica.delete()
        exp_fisica.delete()
        cliente_fisica.delete()

    def test_tareas_no_incluye_expedientes_de_otro_asesor(self):
        r = self.client.get(reverse('extension_api_tareas'), **self._auth(self.otro.profile.api_token))
        self.assertEqual(r.status_code, 200)
        self.assertEqual(r.json()['tareas'], [])

    def test_tareas_solo_modo_extension(self):
        """La extensión solo debe ver tareas creadas para ella (no headless)."""
        # Tarea headless pendiente (no debe aparecer)
        TareaConciliacion.objects.create(
            expediente=self.expediente, usuario=self.asesor,
            estado='pendiente', modo='automatico',
        )
        r = self.client.get(reverse('extension_api_tareas'), **self._auth())
        data = r.json()
        self.assertEqual(len(data['tareas']), 1)  # solo la de modo='extension'
        self.assertEqual(data['tareas'][0]['id'], self.tarea.pk)

    def test_doble_reporte_es_409(self):
        """Una tarea ya reportada no acepta un segundo reporte."""
        r1 = self.client.post(
            reverse('extension_api_reportar', args=[self.tarea.pk]),
            data=json.dumps({'estado': 'completado', 'folio': 'CCL-ONCE'}),
            content_type='application/json',
            **self._auth(),
        )
        self.assertEqual(r1.status_code, 200)

        r2 = self.client.post(
            reverse('extension_api_reportar', args=[self.tarea.pk]),
            data=json.dumps({'estado': 'completado', 'folio': 'CCL-DOS-VECES'}),
            content_type='application/json',
            **self._auth(),
        )
        self.assertEqual(r2.status_code, 409)

        # El folio original se conserva
        self.tarea.refresh_from_db()
        self.assertEqual(self.tarea.folio, 'CCL-ONCE')

    # ─── POST /api/extension/tareas/<id>/reportar/ ────────────────────

    def test_reportar_completado_guarda_folio_y_acuse(self):
        import base64 as b64
        pdf_fake = b'%PDF-1.4 fake acuse content'
        payload = {
            'estado': 'completado',
            'folio': 'CCL-2025-1234',
            'detalle': 'Enviado desde la extensión',
            'acuse_pdf': b64.b64encode(pdf_fake).decode(),
            'acuse_nombre': 'acuse_prueba.pdf',
        }
        r = self.client.post(
            reverse('extension_api_reportar', args=[self.tarea.pk]),
            data=json.dumps(payload),
            content_type='application/json',
            **self._auth(),
        )
        self.assertEqual(r.status_code, 200)
        self.assertTrue(r.json()['ok'])

        self.tarea.refresh_from_db()
        self.expediente.refresh_from_db()
        self.assertEqual(self.tarea.estado, 'completado')
        self.assertEqual(self.tarea.folio, 'CCL-2025-1234')
        self.assertEqual(self.expediente.folio, 'CCL-2025-1234')
        self.assertIsNotNone(self.tarea.completed_at)

        # El acuse se guardó como Documento
        self.assertEqual(self.expediente.documentos.count(), 1)
        doc = self.expediente.documentos.first()
        self.assertEqual(doc.descripcion, 'Acuse de Conciliación - Folio: CCL-2025-1234')
        self.assertTrue(doc.archivo.name.endswith('.pdf'))

    def test_reportar_fallido_guarda_error(self):
        payload = {'estado': 'fallido', 'error': 'El portal rechazó la CURP'}
        r = self.client.post(
            reverse('extension_api_reportar', args=[self.tarea.pk]),
            data=json.dumps(payload),
            content_type='application/json',
            **self._auth(),
        )
        self.assertEqual(r.status_code, 200)
        self.tarea.refresh_from_db()
        self.assertEqual(self.tarea.estado, 'fallido')
        self.assertEqual(self.tarea.error, 'El portal rechazó la CURP')

    def test_reportar_sin_acceso_es_403(self):
        payload = {'estado': 'completado', 'folio': 'X-1'}
        r = self.client.post(
            reverse('extension_api_reportar', args=[self.tarea.pk]),
            data=json.dumps(payload),
            content_type='application/json',
            **self._auth(self.otro.profile.api_token),
        )
        self.assertEqual(r.status_code, 403)

    def test_reportar_sin_token_es_401(self):
        r = self.client.post(
            reverse('extension_api_reportar', args=[self.tarea.pk]),
            data=json.dumps({'estado': 'fallido'}),
            content_type='application/json',
        )
        self.assertEqual(r.status_code, 401)

    def test_reportar_no_exige_csrf_cookie(self):
        """
        La extensión hace POST cross-origin sin cookie CSRF (autentica por token).
        Los endpoints deben estar marcados csrf_exempt; si no, el navegador
        recibe 403 'CSRF cookie not set'.
        """
        self.client.enforce_csrf_checks = True
        payload = {'estado': 'fallido', 'error': 'prueba csrf'}
        r = self.client.post(
            reverse('extension_api_reportar', args=[self.tarea.pk]),
            data=json.dumps(payload),
            content_type='application/json',
            **self._auth(),
        )
        self.assertEqual(r.status_code, 200)
        self.assertTrue(r.json()['ok'])

        # También GET de tareas
        r2 = self.client.get(reverse('extension_api_tareas'), **self._auth())
        self.assertEqual(r2.status_code, 200)

    def test_reportar_con_screenshots_guarda_espejo(self):
        import base64 as b64
        from PIL import Image
        from io import BytesIO

        # PNG pequeño
        buf = BytesIO()
        Image.new('RGB', (10, 10), (255, 0, 0)).save(buf, format='PNG')
        png_b64 = b64.b64encode(buf.getvalue()).decode()

        payload = {'estado': 'completado', 'folio': 'CCL-SHOT-1', 'screenshots': [png_b64]}
        r = self.client.post(
            reverse('extension_api_reportar', args=[self.tarea.pk]),
            data=json.dumps(payload),
            content_type='application/json',
            **self._auth(),
        )
        self.assertEqual(r.status_code, 200)
        self.tarea.refresh_from_db()
        self.assertIn('ext_00.png', self.tarea.screenshots_json)

    # ─── Página de configuración ───────────────────────────────────────

    def test_pagina_config_muestra_token(self):
        self.client.login(username='ext_asesor', password='x')
        r = self.client.get(reverse('extension_config'))
        self.assertEqual(r.status_code, 200)
        self.assertContains(r, self.token)
        self.assertContains(r, 'Extensión de Chrome')

    def test_pagina_config_requiere_login(self):
        r = self.client.get(reverse('extension_config'))
        self.assertEqual(r.status_code, 302)

    def test_regenerar_token_invalida_el_anterior(self):
        self.client.login(username='ext_asesor', password='x')
        viejo = self.asesor.profile.api_token
        r = self.client.post(reverse('extension_regenerar_token'))
        self.assertEqual(r.status_code, 302)
        self.asesor.profile.refresh_from_db()
        self.assertNotEqual(self.asesor.profile.api_token, viejo)
        self.assertTrue(self.asesor.profile.api_token)

    def test_envio_modo_extension_crea_tarea_pendiente_sin_ejecutar(self):
        """Elegir 'extensión' en el confirmar crea la tarea sin lanzar el hilo headless."""
        self.client.login(username='ext_asesor', password='x')
        r = self.client.post(
            reverse('enviar_conciliacion_automation', args=[self.expediente.pk]),
            {'modo': 'extension'},
        )
        # Redirige a la página de configuración de la extensión
        self.assertEqual(r.status_code, 302)
        self.assertIn(reverse('extension_config'), r.url)

        # La nueva tarea quedó pendiente (sin thread)
        tarea_nueva = self.expediente.tareas_conciliacion.filter(modo='extension').order_by('-created_at').first()
        self.assertIsNotNone(tarea_nueva)
        self.assertEqual(tarea_nueva.estado, 'pendiente')


class AvisosObligatoriosTests(TestCase):
    """
    Tests de los avisos obligatorios del admin:

    - El admin publica un aviso desde el dashboard (POST /avisos/crear/)
    - El context processor detecta el aviso no leído (aviso_obligatorio)
    - El usuario lo marca como leído (POST /avisos/<id>/marcar-leido/)
    - No vuelve a aparecer una vez leído
    """

    @classmethod
    def setUpTestData(cls):
        cls.admin = User.objects.create_user(username='avisos_admin', password='x', first_name='Admin', last_name='Avisos')
        cls.admin.profile.rol = 'admin'
        cls.admin.profile.save()

        cls.asesor = User.objects.create_user(username='avisos_asesor', password='x', first_name='Asesor', last_name='Avisos')
        cls.asesor.profile.rol = 'asesor'
        cls.asesor.profile.save()

    def test_crear_aviso_requiere_admin(self):
        self.client.login(username='avisos_asesor', password='x')
        r = self.client.post(reverse('crear_aviso'), {'titulo': 'No permitido'})
        self.assertEqual(r.status_code, 403)
        self.assertEqual(Aviso.objects.count(), 0)

    def test_crear_aviso_desde_dashboard(self):
        self.client.login(username='avisos_admin', password='x')
        r = self.client.post(reverse('crear_aviso'), {
            'titulo': 'Junta general viernes',
            'contenido': 'Todos deben asistir a las 10:00',
            'prioridad': 'alta',
        })
        self.assertEqual(r.status_code, 302)  # redirect al dashboard
        aviso = Aviso.objects.get(titulo='Junta general viernes')
        self.assertEqual(aviso.prioridad, 'alta')
        self.assertTrue(aviso.activo)
        self.assertEqual(aviso.creado_por, self.admin)
        # El creador ya lo marcó como leído (no se auto-muestra)
        self.assertIn(self.admin, aviso.leido_por.all())

    def test_titulo_obligatorio(self):
        self.client.login(username='avisos_admin', password='x')
        r = self.client.post(reverse('crear_aviso'), {'titulo': '   '})
        self.assertEqual(r.status_code, 302)
        self.assertEqual(Aviso.objects.count(), 0)

    def test_context_processor_detecta_aviso_no_leido(self):
        aviso = Aviso.objects.create(
            titulo='Aviso urgente', contenido='Contenido',
            prioridad='alta', activo=True, creado_por=self.admin,
        )
        self.client.login(username='avisos_asesor', password='x')
        r = self.client.get(reverse('dashboard_asesor'))
        self.assertContains(r, 'Aviso urgente')
        self.assertContains(r, 'aviso-obligatorio')  # modal bloqueante presente
        self.assertContains(r, 'Entendido')

    def test_context_processor_no_muestra_aviso_ya_leido(self):
        aviso = Aviso.objects.create(
            titulo='Aviso leído', contenido='Contenido',
            prioridad='media', activo=True, creado_por=self.admin,
        )
        aviso.leido_por.add(self.asesor)
        self.client.login(username='avisos_asesor', password='x')
        r = self.client.get(reverse('dashboard_asesor'))
        self.assertNotContains(r, 'aviso-obligatorio')

    def test_context_processor_no_muestra_aviso_inactivo(self):
        Aviso.objects.create(
            titulo='Aviso inactivo', contenido='x',
            prioridad='media', activo=False, creado_por=self.admin,
        )
        self.client.login(username='avisos_asesor', password='x')
        r = self.client.get(reverse('dashboard_asesor'))
        self.assertNotContains(r, 'aviso-obligatorio')

    def test_context_processor_muestra_el_mas_reciente(self):
        Aviso.objects.create(titulo='Aviso 1', contenido='x', prioridad='baja', activo=True, creado_por=self.admin)
        Aviso.objects.create(titulo='Aviso 2', contenido='x', prioridad='baja', activo=True, creado_por=self.admin)
        self.client.login(username='avisos_asesor', password='x')
        r = self.client.get(reverse('dashboard_asesor'))
        html = r.content.decode()
        # El modal bloqueante (aviso-obligatorio) solo muestra el MÁS RECIENTE
        inicio_modal = html.index('id="aviso-obligatorio"')
        fin_modal = html.index('aviso-entendido')
        modal = html[inicio_modal:fin_modal]
        self.assertIn('Aviso 2', modal)
        self.assertNotIn('Aviso 1', modal)  # la lista semanal sí lo tiene, el modal no

    def test_marcar_leido_quita_el_modal(self):
        aviso = Aviso.objects.create(
            titulo='Para leer', contenido='x',
            prioridad='alta', activo=True, creado_por=self.admin,
        )
        self.client.login(username='avisos_asesor', password='x')

        # Primero aparece
        r1 = self.client.get(reverse('dashboard_asesor'))
        self.assertContains(r1, 'aviso-obligatorio')

        # Marcar leído
        r2 = self.client.post(reverse('marcar_aviso_leido', args=[aviso.pk]))
        self.assertEqual(r2.status_code, 200)
        self.assertTrue(r2.json()['ok'])
        self.assertIn(self.asesor, aviso.leido_por.all())

        # Ya no aparece
        r3 = self.client.get(reverse('dashboard_asesor'))
        self.assertNotContains(r3, 'aviso-obligatorio')

    def test_marcar_leido_requiere_login(self):
        aviso = Aviso.objects.create(titulo='x', activo=True, creado_por=self.admin)
        r = self.client.post(reverse('marcar_aviso_leido', args=[aviso.pk]))
        self.assertEqual(r.status_code, 302)  # redirect al login

    # ─── Fecha de vencimiento ──────────────────────────────────────────

    def test_crear_aviso_con_fecha_vencimiento(self):
        self.client.login(username='avisos_admin', password='x')
        r = self.client.post(reverse('crear_aviso'), {
            'titulo': 'Aviso con vencimiento',
            'contenido': 'Solo hasta el viernes',
            'prioridad': 'alta',
            'fecha_vencimiento': '2026-12-31T18:00',
        })
        self.assertEqual(r.status_code, 302)
        aviso = Aviso.objects.get(titulo='Aviso con vencimiento')
        self.assertIsNotNone(aviso.fecha_vencimiento)
        # La fecha se guarda en UTC (USE_TZ=True); se verifica en hora local
        # de México, que es como la captura el admin y como se muestra.
        fecha_local = timezone.localtime(aviso.fecha_vencimiento)
        self.assertEqual(fecha_local.year, 2026)
        self.assertEqual(fecha_local.month, 12)
        self.assertEqual(fecha_local.day, 31)
        self.assertEqual(fecha_local.hour, 18)

    def test_crear_aviso_fecha_vencimiento_invalida_publica_sin_fecha(self):
        self.client.login(username='avisos_admin', password='x')
        r = self.client.post(reverse('crear_aviso'), {
            'titulo': 'Aviso fecha mala',
            'prioridad': 'media',
            'fecha_vencimiento': 'no-es-una-fecha',
        })
        self.assertEqual(r.status_code, 302)
        aviso = Aviso.objects.get(titulo='Aviso fecha mala')
        self.assertIsNone(aviso.fecha_vencimiento)  # se publica sin vencimiento

    def test_context_processor_no_muestra_aviso_vencido(self):
        Aviso.objects.create(
            titulo='Aviso vencido', contenido='x', prioridad='alta',
            activo=True, creado_por=self.admin,
            fecha_vencimiento=timezone.now() - timedelta(days=1),
        )
        self.client.login(username='avisos_asesor', password='x')
        r = self.client.get(reverse('dashboard_asesor'))
        self.assertNotContains(r, 'aviso-obligatorio')
        self.assertNotContains(r, 'Aviso vencido')  # tampoco en la lista

    def test_context_processor_muestra_aviso_no_vencido(self):
        Aviso.objects.create(
            titulo='Aviso vigente', contenido='x', prioridad='alta',
            activo=True, creado_por=self.admin,
            fecha_vencimiento=timezone.now() + timedelta(days=5),
        )
        self.client.login(username='avisos_asesor', password='x')
        r = self.client.get(reverse('dashboard_asesor'))
        self.assertContains(r, 'aviso-obligatorio')
        self.assertContains(r, 'Aviso vigente')


class AjustesTests(TestCase):
    """Tests del módulo de Ajustes y la descarga del paquete de la extensión."""

    @classmethod
    def setUpTestData(cls):
        cls.asesor = User.objects.create_user(username='ajustes_user', password='x', first_name='A', last_name='Justes')
        cls.asesor.profile.rol = 'asesor'
        cls.asesor.profile.save()

    def test_ajustes_requiere_login(self):
        r = self.client.get(reverse('ajustes'))
        self.assertEqual(r.status_code, 302)

    def test_ajustes_muestra_token_y_seccion_extension(self):
        self.client.login(username='ajustes_user', password='x')
        r = self.client.get(reverse('ajustes'))
        self.assertEqual(r.status_code, 200)
        self.assertContains(r, 'Extensión de Chrome')
        self.assertContains(r, self.asesor.profile.api_token)
        self.assertContains(r, 'Descargar extensión')

    def test_descargar_paquete_zip(self):
        self.client.login(username='ajustes_user', password='x')
        r = self.client.get(reverse('extension_descargar'))
        self.assertEqual(r.status_code, 200)
        self.assertEqual(r['Content-Type'], 'application/zip')
        self.assertIn('conciliacion_bc_extension.zip', r['Content-Disposition'])
        # El zip debe contener el manifest.json
        import zipfile
        from io import BytesIO
        zf = zipfile.ZipFile(BytesIO(r.content))
        nombres = zf.namelist()
        self.assertTrue(any(n.endswith('manifest.json') for n in nombres))
        self.assertTrue(any(n.endswith('content.js') for n in nombres))

    def test_el_icono_de_ajustes_esta_en_base(self):
        """El header debe tener el enlace a Ajustes (ícono de engrane)."""
        self.client.login(username='ajustes_user', password='x')
        r = self.client.get(reverse('dashboard_asesor'))
        self.assertContains(r, reverse('ajustes'))
        self.assertContains(r, 'fa-cog')
